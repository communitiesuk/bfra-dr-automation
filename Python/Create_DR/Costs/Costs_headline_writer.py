# RAP_writer_headline.py
"""
Created on Tuesday 26 August 2025, 11:18:21

Author: Matthew Bandura
"""

import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
import sys
import os

# Add the Utility folder to sys.path
folder_path = os.path.abspath(os.path.join(os.getcwd(), '..', 'Utility'))  # Replace 'folder_name' with the folder's name
sys.path.append(folder_path)

from RAP.RAP_variables import RAP_variable_creator
from Utility.functions import Change_line_in_DR, format_percentage, create_table
import Utility.docx_svg_patch

def RAP_headline_writer(RAP_headline_dict, Estimates_headline_dict, figure_count, dates_variables, DR):
    last_month = dates_variables['last_month']
    this_month = dates_variables['this_month']
    year = dates_variables['year']
    last_month_year = dates_variables['last_month_year']

    RAP_total = RAP_headline_dict['RAP_total']
    RAP_total_line = RAP_headline_dict['RAP_total_line']
    RAP_total_since_oct_23 = RAP_headline_dict['RAP_total_since_oct_23']
    RAP_started_c_no = RAP_headline_dict['RAP_started_c_no']
    RAP_started_c_pct = RAP_headline_dict['RAP_started_c_pct']
    RAP_completed_c_no = RAP_headline_dict['RAP_completed_c_no']
    RAP_completed_c_pct = RAP_headline_dict['RAP_completed_c_pct']

    Estimates_11m_proportion_of_low_estimate = Estimates_headline_dict['Estimates_11m_proportion_of_low_estimate']
    Estimates_11m_proportion_of_high_estimate = Estimates_headline_dict['Estimates_11m_proportion_of_high_estimate']

    # Headline Title
    text = f'Overall remediation'
    paragraph = DR.add_paragraph(text, style = 'Heading 3')

    # Paragraph
    text = f'As at the end of {this_month}, there are {RAP_total} residential buildings 11 metres and over in height identified with unsafe cladding whose remediation progression is being reported on in this release, {RAP_total_line} since the end of {last_month} {last_month_year}. This is an estimated {Estimates_11m_proportion_of_high_estimate}-{Estimates_11m_proportion_of_low_estimate} of all buildings 11 metres and over in height expected to be remediated as part of MHCLG’s remediation programmes.'
    DR.add_paragraph(text, style = 'Normal')

    # Paragraph
    text = f'Since the department first began reporting on all five remediation programmes in October 2023, {RAP_total_since_oct_23} buildings with unsafe cladding are being reported on in this release.'
    DR.add_paragraph(text, style = 'Normal')

    # Paragraph
    text = f'Overall, {RAP_started_c_no} buildings ({RAP_started_c_pct}) have either started or completed remediation works. Of these, {RAP_completed_c_no} buildings ({RAP_completed_c_pct}) have completed remediation works.'
    DR.add_paragraph(text, style = 'Normal')

    # Figure Title
    paragraph = DR.add_paragraph(style = 'Normal')
    text = f'Figure {figure_count}: Of the {RAP_total} buildings identified with unsafe cladding, {RAP_started_c_no} ({RAP_started_c_pct}) have started or completed remediation works, of which {RAP_completed_c_no} ({RAP_completed_c_pct}) have completed remediation works. This includes remediation progress on high rise (18m+) and mid-rise (11-18m) buildings in height.'
    run = paragraph.add_run(text)
    run.bold = True
    figure_count += 1

    # Figure
    DR.add_picture('Q:\\BSP\Automation\\DR Automation\\Excel_inputs\\[ADD FIGURE HERE INSERT]\\Add_figure_here.png', width=Cm(17))

    # Figure caption 1
    DR.add_paragraph('Note: From October 2023 onwards combined remediation progress is shown across the BSF, ACM programme, Cladding Safety Scheme, developer remediation contract and as reported by registered providers of social housing. The total number of buildings identified with unsafe cladding, reported in the overall remediation section of the data release, does not sum to the total number of buildings in each remediation programme, reported in each respective section of the data release. This is due to some buildings appearing in more than one remediation programme.', style = 'Normal')

    # Figure caption 2
    DR.add_paragraph('The increase in the number of reported completions between May and June 2024 is largely due to a change in methodology in reporting social housing sector remediation, which from June 2024 includes social housing buildings that had completed remediation independently of government funding and monitoring schemes before March 2024.', style = 'Normal')

    return figure_count