# Costs_section_writer.py
"""
Created on Monday 22nd December 2025 15:32:21

Author: Matthew Bandura
"""

import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
import sys
import os

# Add the Utility folder to sys.path
folder_path = os.path.abspath(os.path.join(os.getcwd(), '..', 'Utility'))  # Replace 'folder_name' with the folder's name
sys.path.append(folder_path)

from Utility.functions import create_table, add_hyperlink


def Costs_section_writer(Costs_section_dict, Costs_tables, table_count, dates_variables, DR):    
    # Unpacking date variables
    hyperlink_month = dates_variables['hyperlink_month']
    year = dates_variables['year']

    Costs_total_low = Costs_section_dict['Costs_total_low']
    Costs_total_high = Costs_section_dict['Costs_total_high']

    Costs_gvt_low_pct = Costs_section_dict['Costs_gvt_low_pct']
    Costs_gvt_high_pct = Costs_section_dict['Costs_gvt_high_pct']
    Costs_gvt_low = Costs_section_dict['Costs_gvt_low']
    Costs_gvt_high = Costs_section_dict['Costs_gvt_high']


    Costs_other_high = Costs_section_dict['Costs_other_high']
    Costs_other_low = Costs_section_dict['Costs_other_low']
    Costs_other_low_pct = Costs_section_dict['Costs_other_low_pct']
    Costs_other_high_pct = Costs_section_dict['Costs_other_high_pct']

    Costs_total_available = Costs_section_dict['Costs_total_available']
    Costs_exchequer = Costs_section_dict['Costs_exchequer']
    Costs_developer = Costs_section_dict['Costs_developer']
    Costs_BSL = Costs_section_dict['Costs_BSL']



    Costs_estimates_table = Costs_tables['Costs_estimates_table']
    Costs_proportions_table = Costs_tables['Costs_proportions_table']

    # Section Title 
    paragraph = DR.add_paragraph('Estimated cost of external wall system\n remediation of 11m+ residential buildings in\n England', style = 'Heading 2')

    # Intro
    paragraph = DR.add_paragraph('Key statistics:', style = 'Heading 3')

    # Paragraph 1
    paragraph = DR.add_paragraph(style = 'Normal')
    paragraph.add_run(f'As of March 2025, MHCLG estimates that it will cost between £{Costs_total_low}bn and £{Costs_total_high}bn to remediate residential buildings, containing dwellings,')
    paragraph.add_run(' 11 metres and above in height that have or had unsafe cladding in England.')

    #Paragraph 2
    paragraph = DR.add_paragraph(f'Of this total estimated capital cost:', style = 'Normal')

    #Bullet Point 1
    paragraph = DR.add_paragraph(f'An estimated £{Costs_gvt_high}bn - £{Costs_gvt_low}bn will have been funded by government once all works have been completed ({Costs_gvt_low_pct} - {Costs_gvt_high_pct} of the total capital costs;', style = 'List Bullet')

    #Bullet Point 2
    paragraph = DR.add_paragraph(f'An estimated £{Costs_other_high}bn - £{Costs_other_low}bn will have been funded by government once all works have been completed ({Costs_other_low_pct} - {Costs_other_high_pct} of the total capital costs. This includes developers, registered providers of social housing and other actors.', style = 'List Bullet')

    # Paragraph 3
    paragraph = DR.add_paragraph(style = 'Normal')
    paragraph.add_run('These costs relate to buildings which will be or have been remediated or mitigated as part of the department’s five remediation programmes. The costs also include an estimated number of 11m+ buildings which are assumed to be remediated privately,')
    paragraph.add_run(' outside of one of the department’s remediation programmes. They include the estimated cost of works to-date as well as future works. The estimates account for the change in funding eligibility criteria where social landlords will now receive equal access to government remediation funding as private landlords.')
    paragraph.add_run(' They exclude costs which fall outside the scope of remediating buildings 11 metres and above in height that have or had unsafe cladding in England.')


    # Paragraph 4
    paragraph = DR.add_paragraph('These estimates represent our current best estimates and will be subject to change as more information becomes available. They update estimates previously published in the National Audit Office (NAO) report ', style = 'Normal')  
    add_hyperlink(paragraph, '“Dangerous cladding: the government’s remediation portfolio”', 'https://www.nao.org.uk/wp-content/uploads/2024/10/dangerous-cladding-the-governments-remediation-portfolio.pdf')
    paragraph.add_run(f' in November 2024.')

    #Paragraph 5
    paragraph = DR.add_paragraph('More information on the methodology is available in the ')
    add_hyperlink(paragraph, 'Technical Note', f'https://www.gov.uk/government/publications/building-safety-remediation-monthly-data-release-{hyperlink_month}-{year}/')
    paragraph.add_run(f'.')


    # Table 3 caption
    text = f'Table {table_count}: Estimated capital cost of external wall system works to residential buildings over 11 metres with unsafe cladding to be, and which have been, remediated or mitigated, by funding source, England, March 2025, £bn (nominal terms, nearest £0.1bn).'
    paragraph = DR.add_paragraph(style = 'Normal')
    run = paragraph.add_run(text)
    run.bold = True
    table_count += 1

    # Table 3
    table_data = Costs_estimates_table
    table_widths = [Cm(6.5), Cm(2.65), Cm(2.65), Cm(2.75)]
    table_heights = [Cm(1.12), Cm(0.55), Cm(1.13), Cm(0.55), Cm(1.13), Cm(1.13), Cm(0.55)]
    create_table(DR, table_data, table_widths, table_heights)

    # Table 4 caption
    text = f'Table {table_count}: Estimated proportion of capital costs of external wall system works to buildings over 11 metres with unsafe cladding expected to be, and which have been, remediated or mitigated, by funding source, England, March 2025.'
    paragraph = DR.add_paragraph(style = 'Normal')
    run = paragraph.add_run(text)
    run.bold = True
    table_count += 1

    # Table 4
    table_data = Costs_proportions_table
    table_widths = [Cm(6.5), Cm(2.65), Cm(2.65), Cm(2.75)]
    table_heights = [Cm(1.12), Cm(0.55), Cm(1.13), Cm(0.55), Cm(1.13), Cm(1.13), Cm(0.55)]
    create_table(DR, table_data, table_widths, table_heights)

    #Second part intro
    paragraph = DR.add_paragraph('Funding sources available to MHCLG for remediation:', style = 'Heading 3')

    #Paragraph 6
    paragraph = DR.add_paragraph(style = 'Normal')
    paragraph.add_run(f'It is estimated that under current plans, MHCLG will have £{Costs_total_available}bn available to fund the remediation of external wall system defects.')

    #Paragraph 7
    paragraph = DR.add_paragraph(style = 'Normal')
    paragraph.add_run(f'This includes:')

    #Bullet Point 1
    paragraph = DR.add_paragraph(f'a £{Costs_exchequer}bn Exchequer contribution;', style = 'List Bullet')

    #Bullet Point 2
    paragraph = DR.add_paragraph(f'an estimated £{Costs_developer}bn in refunds from developers;', style = 'List Bullet')

    #Bullet Point 3
    paragraph = DR.add_paragraph(f'a £{Costs_BSL}bn Building Safety Levy revenue target.', style = 'List Bullet')

    #Paragraph 8
    paragraph = DR.add_paragraph(style = 'Normal')
    paragraph.add_run(f'As cost estimates continue to be updated, the Levy revenue target will be reviewed and adjusted such that funding available for remediation aligns to government expenditure.')



    return table_count