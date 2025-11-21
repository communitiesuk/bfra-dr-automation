# functions.py
"""
Created on Monday 16 December 2024, 14:58:44

Author: Harry Simmons
"""

import inflect
import os
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_ROW_HEIGHT_RULE

# Function to format percentages and make 0% <1% and 100% <100%
def format_percentage(value):
    if 0 < value < 0.005:
        return '<1%'
    elif 0.995 <= value < 1:
        return '<100%'
    else:
        return f"{round(value * 100)}%"



# Function to set cell borders to be used to add border lines
def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        if edge in kwargs:
            element = tcBorders.find(qn(f'w:{edge}'))
            if element is None:
                element = OxmlElement(f'w:{edge}')
                tcBorders.append(element)
            for key in ['sz', 'val', 'color', 'space', 'shadow']:
                if key in kwargs[edge]:
                    element.set(qn(f'w:{key}'), str(kwargs[edge][key]))

# Function to make rows bold used for tables
def make_text_bold(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

def Change_line_in_DR(value):
    p = inflect.engine()
    if value >= 10 or value <= -10:
        return f"an increase of {value}" if value > 0 else f"a decrease of {abs(value)}"
    elif value > 0:
        return f"an increase of {p.number_to_words(value)}"
    elif value < 0:
        return f"a decrease of {p.number_to_words(abs(value))}"
    else:
        return "no change"

def convert_number(num):
    p = inflect.engine()
    if num < 10:
        return p.number_to_words(num)
    else:
        return str(num)

def create_table(DR, table_data, table_widths, table_heights):
    # Create table
    table = DR.add_table(rows=table_data.shape[0] + 1, cols=table_data.shape[1])

    # Add column headings to the table
    hdr_cells = table.rows[0].cells
    for col_idx in range(table_data.shape[1]):
        hdr_cells[col_idx].text = table_data.columns[col_idx]

    # Add data to the table
    for row_idx in range(table_data.shape[0]):
        row = table.rows[row_idx + 1]
        for col_idx in range(table_data.shape[1]):
            cell = row.cells[col_idx]
            cell.text = str(table_data.iloc[row_idx, col_idx])

    # Set column widths
    for row in table.rows:
        for col_idx, width in enumerate(table_widths):
            row.cells[col_idx].width = width

    # Set row heights
    for row_idx, row in enumerate(table.rows):
        if row_idx < len(table_heights):
            row.height = table_heights[row_idx]
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    # Set bottom black borders for all cells
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, bottom={"sz": 6, "val": "single", "color": "000000"})

    # Make the text in the top row bold
    for cell in table.rows[0].cells:
        make_text_bold(cell)

    # Make the text in the last row bold
    for cell in table.rows[-1].cells:
        make_text_bold(cell)

    return table

def more_or_fewer(value):
    p = inflect.engine()
    # Case when value is 0
    if value == 0:
        return "no additional"
    # If the absolute value is greater than or equal to 11, format with commas
    if abs(value) >= 10:
        formatted_value = format(abs(value), ',')
        return f'{formatted_value} more' if value > 0 else f'{formatted_value} fewer'
    # For smaller values, use words
    if value > 0:
        return f'{p.number_to_words(value)} more'
    elif value < 0:
        return f'{p.number_to_words(abs(value))} fewer'

def number_or_none(value):
    p = inflect.engine()
    if value >= 10:
        return f'{value}'
    elif value > 0:
        return f'{p.number_to_words(value)}'
    else:
        return 'none'

def create_bullet_points_forecast(doc, forecast, enforced, end_period, count):
    if forecast != 'zero':
        count[0] += 1
        if count[0] == 1:
            if forecast == 'one':
                bullet_point = f"One building is forecast to start works by the end of {end_period}"
            else:
                bullet_point = f"{forecast.capitalize()} buildings are forecast to start works by the end of {end_period}"
        elif count[0] == 2:
            if forecast == 'one':
                bullet_point = f"An additional building is forecast to start works by the end of {end_period}"
            else:
                bullet_point = f"An additional {forecast} buildings are forecast to start works by the end of {end_period}"
        else:
            if forecast == 'one':
                bullet_point = f"A further building is forecast to start works by the end of {end_period}"
            else:
                bullet_point = f"{forecast.capitalize()} further buildings are forecast to start works by the end of {end_period}"
        
        if enforced != 'zero':
            if forecast == 'one' and enforced == 'one':
                bullet_point += f", and has had local authority enforcement action taken against it."
            elif enforced == 'one':
                bullet_point += f', and one of these have had local authority enforcement action taken against it.'
            elif forecast == enforced:
                bullet_point += f', all of which have had local authority enforcement action taken against them.'
            else:
                bullet_point += f", and {enforced} of these have had local authority enforcement action taken against them."
        else:
            bullet_point += "."
        
        doc.add_paragraph(bullet_point, style = 'List Bullet')

    return count

def chop_df(df, top_rows_to_drop, bottom_rows_to_drop):
    # Drop rows from 0 to num_rows_to_drop - 1
    df = df.drop(range(top_rows_to_drop))

    # Reset the index
    df = df.reset_index(drop=True)

    # Promote the next row to be the headers
    df.columns = df.iloc[0]
    df = df.drop(0).reset_index(drop=True)

    # Drop bottom rows if bottom_rows_to_drop is not None
    if bottom_rows_to_drop is not None:
        df = df.iloc[:bottom_rows_to_drop]

    return df

def get_excel_path(folder_path):
    files = os.listdir(folder_path)
    excel_files = [file for file in files if file.endswith('.xlsx') or file.endswith('.xls') or file.endswith('.xlsm')]
    if len(excel_files) == 1:
        excel_path = os.path.join(folder_path, excel_files[0])
        return excel_path
    else:
        raise FileNotFoundError("No Excel file found in the folder.")

def add_hyperlink(paragraph, text, url):
    # Create the w:hyperlink tag and add needed values
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a w:r element and a new w:rPr element
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Add formatting to the hyperlink (blue color and underline)
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')  # Blue color
    rPr.append(color)

    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')  # Underline
    rPr.append(underline)

    # Join all the xml elements together and add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Append the hyperlink and return it
    paragraph._element.append(hyperlink)


#function to create the paths to important directories on the user's local machine
def create_paths():

    # Build the path to the graph output folder
    figure_path = r'Q:\BSP\Automation\DR Automation\DR_outputs\DR_graphs'

    placeholder_path = os.path.join(r'Q:\BSP\Automation\DR Automation\Excel_inputs\[ADD FIGURE HERE INSERT]\Add_figure_here.png')

    previous_MI_folder = (r'Q:\BSP\Automation\DR Automation\Excel_inputs\[PUT MI TABLES HERE]\[LAST MONTHS MI TABLES]')
    previous_tables_path = get_excel_path(previous_MI_folder)

    previous_social_folder = (r'Q:\BSP\Automation\DR Automation\Excel_inputs\[PUT MI TABLES HERE]\[LAST MONTHS MI TABLES]\[LAST SOCIAL RELEASE MI TABLES]')
    previous_social_path = get_excel_path(previous_social_folder)

    MI_folder = r"Q:\BSP\Automation\DR Automation\Excel_inputs\[PUT MI TABLES HERE]"
    MI_tables_path = get_excel_path(MI_folder)

    additional_folder = r'Q:\BSP\Automation\DR Automation\Excel_inputs\[PUT ADDITIONAL DR STATS HERE]'
    additional_tables_path = get_excel_path(additional_folder)

    save_path = r'Q:\BSP\Automation\DR Automation\DR_outputs\Auto_DR'

    paths = {'figure_path' : figure_path,
             'previous_tables_path' : previous_tables_path,
             'MI_tables_path' : MI_tables_path,
             'additional_tables_path' : additional_tables_path,
             'save_path' : save_path, 
             'placeholder_path' : placeholder_path,
             'previous_social_path' : previous_social_path
    }
    return paths




