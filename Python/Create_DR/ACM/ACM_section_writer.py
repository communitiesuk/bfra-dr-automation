 # ACM_writer_section.py
"""
Created on Monday 16 December 2024, 15:57:33

Author: Harry Simmons
"""

from docx.shared import Cm 
from docx.enum.text import WD_COLOR_INDEX
import os


from Utility.functions import create_table, create_bullet_points_forecast 

# ACM section writer
def ACM_section_writer(ACM_section_dict, ACM_tables, figure_count, table_count, dates_variables, paths_variables, DR):
    
    placeholder_path = paths_variables['placeholder_path']

    cutoff = dates_variables['cutoff']
    last_month = dates_variables['last_month']
    year = dates_variables['year']
    next_year = dates_variables['next_year']
    end_year_word = dates_variables['end_year_word']
    end_quarter_word = dates_variables['end_quarter_word']

    ACM_remediation_table = ACM_tables['ACM_remediation_table']
    ACM_enforcement_table = ACM_tables['ACM_enforcement_table']

    ACM_yearly_identified = ACM_tables['ACM_yearly_identified']
    ACM_yearly_identified_line = ACM_tables['ACM_yearly_identified_line']

    ACM_started_c_pct = ACM_section_dict['ACM_started_c_pct']
    ACM_started_c_no = ACM_section_dict['ACM_started_c_no']
    ACM_started_line = ACM_section_dict['ACM_started_line']

    ACM_total = ACM_section_dict['ACM_total']
    ACM_total_line = ACM_section_dict['ACM_total_line']
    
    ACM_removed_c_pct = ACM_section_dict['ACM_removed_c_pct']


    ACM_signoff_c_pct = ACM_section_dict['ACM_signoff_c_pct']
    ACM_signoff_c_no = ACM_section_dict['ACM_signoff_c_no']
    ACM_signoff_line = ACM_section_dict['ACM_signoff_line']

    ACM_complete_c_no = ACM_section_dict['ACM_complete_c_no']
    ACM_complete_c_pct = ACM_section_dict['ACM_complete_c_pct']
    ACM_complete_line = ACM_section_dict['ACM_complete_line']

    
    ACM_removed_c_no = ACM_section_dict['ACM_removed_c_no']
    ACM_removed_line = ACM_section_dict['ACM_removed_line']

    ACM_enforcement_total = ACM_section_dict['ACM_enforcement_total']
    ACM_enforcement_pct = ACM_section_dict['ACM_enforcement_pct']
    ACM_enforcement_line = ACM_section_dict['ACM_enforcement_line']
    ACM_enforcement_total_nvt = ACM_section_dict['ACM_enforcement_total_nvt']


    ACM_social_started_c_pct = ACM_section_dict['ACM_social_started_c_pct']
    ACM_social_total = ACM_section_dict['ACM_social_total']

    ACM_private_started_c_pct = ACM_section_dict['ACM_private_started_c_pct']
    ACM_private_total = ACM_section_dict['ACM_private_total']

    ACM_enforcement_forecast_end_quarter_word = ACM_section_dict['ACM_enforcement_forecast_end_quarter_word']
    ACM_enforcement_forecast_this_year_word = ACM_section_dict['ACM_enforcement_forecast_this_year_word']
    ACM_enforcement_forecast_next_year_word = ACM_section_dict['ACM_enforcement_forecast_next_year_word']
    ACM_enforcement_enforced_end_quarter_word = ACM_section_dict['ACM_enforcement_enforced_end_quarter_word']
    ACM_enforcement_enforced_this_year_word = ACM_section_dict['ACM_enforcement_enforced_this_year_word']
    ACM_enforcement_enforced_next_year_word = ACM_section_dict['ACM_enforcement_enforced_next_year_word']
    ACM_enforcement_enforced_no_forecast_word = ACM_section_dict['ACM_enforcement_enforced_no_forecast_word']
    ACM_enforcement_not_enforced_no_forecast_word = ACM_section_dict['ACM_enforcement_not_enforced_no_forecast_word']

    ACM_quarter_progress = ACM_section_dict['ACM_quarter_progress']
  
    ACM_completed_dwellings_low = ACM_section_dict['ACM_completed_dwellings_low']
    ACM_completed_dwellings_high = ACM_section_dict['ACM_completed_dwellings_high']
    
    ACM_yet_to_dwellings_low = ACM_section_dict['ACM_yet_to_dwellings_low']
    ACM_yet_to_dwellings_high = ACM_section_dict['ACM_yet_to_dwellings_high']

    # Section Title
    paragraph = DR.add_paragraph('ACM Remediation', style = 'Heading 2')

    # Paragraph 
    text = f'Information in this section is correct as at {cutoff} and shows a monthly update from the previous publication.'
    paragraph = DR.add_paragraph(text, style = 'Normal')

    # Figure Title
 
    text = f'Figure {figure_count}: {ACM_started_c_pct} of the {ACM_total} identified ACM clad high-rise buildings have started or completed remediation, with {ACM_removed_c_pct} having had their ACM cladding removed and {ACM_signoff_c_pct} having completed remediation, including those awaiting building control sign-off.'
    paragraph = DR.add_paragraph(style = 'Normal')
    run = paragraph.add_run(text)
    run.bold = True
    

    # Figure
    DR.add_picture(os.path.join(paths_variables['figure_path'], f'Figure{figure_count}.svg'), width=Cm(17))
    figure_count += 1

    # Table caption
    text = f'Table {table_count}: Remediation status of buildings with ACM cladding systems unlikely to meet Building Regulations, {cutoff}'
    paragraph = DR.add_paragraph(style = 'Normal')
    run = paragraph.add_run(text)
    run.bold = True
    table_count += 1

    # Table
    table_data = ACM_remediation_table
    table_widths = [Cm(6.5), Cm(2.65), Cm(2.65), Cm(2.75), Cm(2.75)]
    table_heights = [Cm(1.12), Cm(0.55), Cm(1.13), Cm(1.13), Cm(0.55), Cm(0.55), Cm(0.55), Cm(0.55), Cm(0.55)]
    create_table(DR, table_data, table_widths, table_heights)

    # Heading
    paragraph = DR.add_paragraph('ACM Remediation: key statistics', style = 'Heading 3')

    # Paragraph 
    text = f'As of {cutoff}, the department has identified {ACM_total} high-rise residential and publicly owned buildings identified with ACM cladding systems unlikely to meet Building Regulations, {ACM_total_line} since the end of {last_month}.'
    DR.add_paragraph(text, style = 'Normal')

    # Bullet point
    text = f'{ACM_signoff_c_no} buildings ({ACM_signoff_c_pct} of all buildings) have completed ACM remediation – {ACM_signoff_line} since the end of {last_month}. '
    text += f'Of these, {ACM_complete_c_no} buildings ({ACM_complete_c_pct} of all buildings) have received building control sign off – {ACM_complete_line} since the end of {last_month}.'
    DR.add_paragraph(text, style = 'List Bullet')

    # Bullet point
    text = f'{ACM_started_c_no} buildings ({ACM_started_c_pct} of all buildings) have started or completed ACM remediation – {ACM_started_line} since the end of {last_month}.'
    text += f' Of these, {ACM_removed_c_no} buildings ({ACM_removed_c_pct} of all buildings) have removed ACM cladding – {ACM_removed_line} since the end of {last_month}.'
    DR.add_paragraph(text, style = 'List Bullet')

    # Paragraph 
    text = f'There are an estimated {ACM_completed_dwellings_low}-{ACM_completed_dwellings_high} dwellings in private and social sector buildings that have completed remediation, '
    text += f'and a further {ACM_yet_to_dwellings_low}-{ACM_yet_to_dwellings_high} dwellings in occupied private and social sector buildings that have yet to be remediated.'
    DR.add_paragraph(text, style = 'Normal')

    # Heading
    paragraph = DR.add_paragraph('Driving ACM remediation forward', style = 'Heading 3')

    # Paragraph
    text = f'There are {ACM_enforcement_total} buildings yet to start ACM remediation ({ACM_enforcement_pct} of all buildings) - {ACM_enforcement_line} since the end of {last_month}. One building is vacant so does not pose a risk to resident safety.'
    DR.add_paragraph(text, style = 'Normal')

    # Table caption
    text = f'Table {table_count}: Enforcement action and forecast start dates for occupied high-rise buildings yet to start ACM remediation, {cutoff}'
    paragraph = DR.add_paragraph(style = 'Normal')
    run = paragraph.add_run(text)
    run.bold = True
    table_count += 1

    # Table
    table_data = ACM_enforcement_table
    table_widths = [Cm(6.5), Cm(2.4), Cm(2.65), Cm(2.75), Cm(3.4)]
    table_heights = [Cm(3.33), Cm(0.55)]
    table = create_table(DR, table_data, table_widths, table_heights)

    # Paragraph
    text = f'Of the {ACM_enforcement_total_nvt} high-rise occupied buildings yet to start ACM remediation:'
    DR.add_paragraph(text, style = 'Normal')

    count = [0]

    # Bullet point
    create_bullet_points_forecast(DR, ACM_enforcement_forecast_end_quarter_word, ACM_enforcement_enforced_end_quarter_word, end_quarter_word, count)

    # Bullet point
    create_bullet_points_forecast(DR,ACM_enforcement_forecast_this_year_word, ACM_enforcement_enforced_this_year_word, end_year_word, count)

    # Bullet point
    create_bullet_points_forecast(DR, ACM_enforcement_forecast_next_year_word, ACM_enforcement_enforced_next_year_word, next_year, count)

    # Bullet point
    if ACM_enforcement_enforced_no_forecast_word != 'zero':
        if ACM_enforcement_enforced_no_forecast_word == 'one':
            text = 'One building without a forecast start date has had local authority enforcement action taken against it.'
        else:
            text = f'{ACM_enforcement_enforced_no_forecast_word.capitalize()} buildings without a forecast start date have had local authority enforcement action taken against them.'
        DR.add_paragraph(text, style = 'List Bullet')

    # Bullet point
    if ACM_enforcement_not_enforced_no_forecast_word != 'zero':
        if ACM_enforcement_not_enforced_no_forecast_word == 'one':
            text = 'The remaining building has been determined as in scope of the ACM monitoring programme in 2025, and the department continues to engage with building owners to ensure their remediation is progressed.'
        else:
            text = f'The remaining {ACM_enforcement_not_enforced_no_forecast_word} buildings were determined as in scope of the ACM monitoring programme in 2025. The department continues to engage with building owners to ensure their remediation is progressed.'
        DR.add_paragraph(text, style = 'List Bullet')


    # Paragraph
    DR.add_paragraph('These forecast estimates are based on information provided by building owners and agents and may change as further information is received. These estimates can also change as a result of buildings being newly identified. The department continues to engage with building owners to start remediation works on site as soon as possible, and will continue to support local authorities and fire and rescue services in the use of their enforcement powers.', style = 'Normal')

    # Figure Title
    paragraph = DR.add_paragraph(style = 'Normal')
    text = f'Figure {figure_count}: {ACM_quarter_progress} of buildings are forecast to have started or completed ACM remediation works by the end of {end_quarter_word}.'
    run = paragraph.add_run(text)
    run.bold = True
    

    # Figure 
    DR.add_picture(placeholder_path, width=Cm(17))
    figure_count += 1

    # Heading
    paragraph = DR.add_paragraph('ACM remediation progress by year of identification', style = 'Heading 3')

    # Figure Title
    paragraph = DR.add_paragraph(style = 'Normal')
    text= f'Figure {figure_count}: '
    run = paragraph.add_run(text)
    run.bold = True
    run = paragraph.add_run('[ADD STATISTIC]')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    text = f' of buildings identified at 31 December 2024 have started or completed remediation compared to {ACM_started_c_pct} of all buildings in the programme.'
    run = paragraph.add_run(text)
    run.bold = True
    

    # Figure
    DR.add_picture(placeholder_path, width=Cm(17))
    figure_count += 1

    # Paragraph
    text = f"Since 31 December 2021, {ACM_yearly_identified_line['Number of buildings identified'].sum()} further high-rise residential buildings have been identified with ACM cladding systems unlikely to meet Building Regulations and have moved into scope of the Building Safety Programme. Of these,"
    for index, row in ACM_yearly_identified_line.iterrows():
        year = row['Year of identification']
        count = row['Number of buildings identified']
        words = row['Words']
        if count > 0:
            text += f" {words} buildings were identified in {year},"
    text = text.rstrip(',')
    parts = text.rsplit(',', 1)
    text = ' and'.join(parts)
    text += '.'
    DR.add_paragraph(text, style = 'Normal')

    # Table caption
    text = f'Table {table_count}: Buildings with unsafe ACM cladding by year of identification, {cutoff}'
    paragraph = DR.add_paragraph(style = 'Normal')
    run = paragraph.add_run(text)
    run.bold = True
    table_count += 1

    # Table
    table_data = ACM_yearly_identified
    table_widths = [Cm(5)]
    table_heights = [Cm(0.55), Cm(0.55), Cm(0.55),Cm(0.55), Cm(0.55), Cm(0.55), Cm(0.55), Cm(0.55), Cm(0.55), Cm(0.55), Cm(0.55)]
    create_table(DR, table_data, table_widths, table_heights)

    # Heading
    paragraph = DR.add_paragraph('ACM remediation by sector', style = 'Heading 3')

    # Paragraph

    paragraph = DR.add_paragraph(style = 'Normal')
    text = f'Figure {figure_count}: {ACM_social_started_c_pct} of the {ACM_social_total} social sector residential buildings in the ACM programme have started or completed remediation, compared to {ACM_private_started_c_pct} of the {ACM_private_total} private sector residential buildings.'
    run = paragraph.add_run(text)
    run.bold = True

    # Figure
    DR.add_picture(os.path.join(paths_variables['figure_path'], f'Figure{figure_count}.svg'), width=Cm(17))
    figure_count += 1

    return figure_count, table_count