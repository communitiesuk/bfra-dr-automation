# Estimates_writer_section.py
"""
Created on Monday 16 December 2024, 15:57:33

Author: Harry Simmons
"""

from docx.shared import Cm
import sys
import os

# Add the Utility folder to sys.path
folder_path = os.path.abspath(os.path.join(os.getcwd(), '..', 'Utility'))  # Replace 'folder_name' with the folder's name
sys.path.append(folder_path)

from Utility.functions import create_table, add_hyperlink
import Utility.docx_svg_patch

# Estimates section writer
def Estimates_section_writer(Estimates_section_dict, Estimates_tables, table_count, dates_variables, DR):
    cutoff = dates_variables['cutoff']
    this_month = dates_variables['this_month']
    hyperlink_month = dates_variables['hyperlink_month']
    
    Estimates_tables_number_of_buildings = Estimates_tables['Estimates_tables_number_of_buildings']
    Estimates_tables_proportion_of_buildings = Estimates_tables['Estimates_tables_proportion_of_buildings']
    Estimates_11m_proportion_of_low_estimate = Estimates_section_dict['Estimates_11m_proportion_of_low_estimate']
    Estimates_11m_proportion_of_high_estimate = Estimates_section_dict['Estimates_11m_proportion_of_high_estimate']

    # Section Title
    paragraph = DR.add_paragraph('Estimated number of buildings with unsafe cladding', style = 'Heading 2')

    # Paragraph 
    paragraph = DR.add_paragraph('Key statistics:', style = 'Heading 3')

    # Paragraph 
    text = f'MHCLG’s best estimates, as of January 2025, are that there are between 5,900 and 9,000 residential buildings, containing dwellings, 11 metres and over in height that have or had unsafe cladding requiring work in England. This equates to an estimated 12-13% of the residential building stock, containing dwellings, over 11 metres in England.'
    DR.add_paragraph(text, style = 'Normal')

    # Paragraph 
    text = f'Of these:'
    DR.add_paragraph(text, style = 'Normal')

    # Bullet point
    text = f'An estimated 2,900 – 5,800 buildings are 11-18m (7-10% of the estimated 39,000 - 59,000 11-18m buildings).'
    DR.add_paragraph(text, style = 'List Bullet')

    # Bullet point
    text = f'An estimated 2,900 - 3,200 buildings are 18m+ (24-27% of the estimated 12,000 18m+ stock).'
    DR.add_paragraph(text, style = 'List Bullet')

    # Paragraph 
    text = f'Of the 11-18m buildings requiring work:'
    DR.add_paragraph(text, style = 'Normal')

    # Bullet point
    text = f'78% are 11-<14m buildings (7-8% of the estimated 35,000 - 53,000 11-<14m buildings).'
    DR.add_paragraph(text, style = 'List Bullet')

    # Bullet point
    text = f'22% are 14-18m buildings (15-19% of the estimated 4,000 – 7,000 14-18m buildings).'
    DR.add_paragraph(text, style = 'List Bullet')

    # Paragraph
    text = f'The estimates of the number of 11-18m buildings in England and the subset of those with unsafe cladding, as of January 2025, replace the estimates first published in 2021 and in 2022. New estimates use the recently released Ordnance Survey National Geographic Database (OS NGD) data which was not available when the first estimates were calculated. More information on the methodology is available in the'
    paragraph = DR.add_paragraph(text, style = 'Normal')
    add_hyperlink(paragraph, 'Technical Note', f'https://www.gov.uk/government/publications/building-safety-remediation-monthly-data-release-{hyperlink_month}')
    paragraph.add_run('.')

    # Heading
    paragraph = DR.add_paragraph('Estimates of buildings to be remediated in MHCLG’s remediation programmes:', style = 'Heading 3')

    # Paragraph
    text = f'MHCLG estimate there to be between 5,700 and 8,600 residential buildings, 11 metres over in height that have or had unsafe cladding that will be or have been remediated or mitigated as part of the department’s five remediation programmes, and therefore will be or are already monitored in the Building Safety Remediation monthly data release.'
    DR.add_paragraph(text, style = 'Normal')

    # Paragraph 
    text = f'These figures differ from those in the Key Statistics section above because:'
    DR.add_paragraph(text, style = 'Normal')

    # Bullet point
    text = f'18m+ hotels, student accommodation or public buildings identified with unsafe ACM cladding are included in the estimates in this section, because these buildings are remediated through the ACM programme, but are excluded from the Key Statistics section above because the buildings do not contain dwellings. At the time the estimate was made there were 98 of these buildings.'
    DR.add_paragraph(text, style = 'List Bullet')

    # Bullet point
    text = f'The estimated number of 11m+ buildings which are assumed to be remediated privately, outside of one of the department’s remediation programmes, are included in the Key Statistics section above but excluded from the estimates in this section.'
    DR.add_paragraph(text, style = 'List Bullet')

    # Paragraph 
    text = f'Of these:'
    DR.add_paragraph(text, style = 'Normal')

    # Bullet point
    text = f'An estimated 2,800 – 5,400 buildings are 11-18m.'
    DR.add_paragraph(text, style = 'List Bullet')

    # Bullet point
    text = f'An estimated 3,000 – 3,200 buildings are 18m+.'
    DR.add_paragraph(text, style = 'List Bullet')

    # Table caption
    text = f'Table {table_count}: Estimated number of residential buildings with unsafe cladding expected to be remediated or mitigated as part of MHCLG’s remediation programmes, over 11 metres with unsafe cladding, by height, England, January 2025. These figures include non dwelling 18m+ buildings with unsafe cladding, and excludes buildings assumed to be remediated outside of a remediation programme. Totals do not sum due to rounding.'
    paragraph = DR.add_paragraph(style = 'Normal')
    run = paragraph.add_run(text)
    run.bold = True
    table_count += 1

    # Table
    table_data = Estimates_tables_number_of_buildings
    table_widths = [Cm(6.5), Cm(6.5), Cm(6.5)]
    table_heights = [Cm(1.1), Cm(1.1), Cm(1.1)]
    create_table(DR, table_data, table_widths, table_heights)

    # Table caption
    text = f'Table {table_count}: Estimated proportion of buildings with unsafe cladding expected to be remediated or mitigated in a government remediation programme that are currently monitored in a government remediation programme, {cutoff}. Proportions are calculated from unrounded estimates so may not be derivable from rounded estimates in the data release.'
    paragraph = DR.add_paragraph(style = 'Normal')
    run = paragraph.add_run(text)
    run.bold = True
    table_count += 1

    # Table
    table_data = Estimates_tables_proportion_of_buildings
    table_widths = [Cm(6.5), Cm(6.5), Cm(6.5), Cm(6.5)]
    table_heights = [Cm(1.1), Cm(1.1), Cm(1.1), Cm(1.1)]
    create_table(DR, table_data, table_widths, table_heights)

    # Paragraph 
    text = f'As at the end of {this_month}, MHCLG is monitoring the remediation progress of an estimated {Estimates_11m_proportion_of_high_estimate}-{Estimates_11m_proportion_of_low_estimate} of residential buildings with unsafe cladding expected to be remediated or mitigated in the department’s remediation programmes. More information on the remediation progress of buildings MHCLG are currently monitoring can be found in the '
    paragraph = DR.add_paragraph(text, style = 'Normal')
    add_hyperlink(paragraph, 'Overall Remediation section', f'https://www.gov.uk/government/publications/building-safety-remediation-monthly-data-release-{hyperlink_month}/building-safety-remediation-monthly-data-release-{hyperlink_month}#overall-remediation-progress')                 
    paragraph.add_run(' of the data release.')

    # Paragraph 
    text = f'These figures include buildings that have already completed remediation or mitigation of unsafe cladding.'
    DR.add_paragraph(text, style = 'Normal')

    return table_count