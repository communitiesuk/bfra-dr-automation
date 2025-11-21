# Developer_writer_section.py
"""
Created on Wednesday 19 February 2025, 14:39:56

Author: Harry Simmons
"""

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_COLOR_INDEX
import sys
import os

# Add the Utility folder to sys.path
folder_path = os.path.abspath(os.path.join(os.getcwd(), '..', 'Utility'))  # Replace 'folder_name' with the folder's name
sys.path.append(folder_path)

from Utility.functions import create_table, add_hyperlink

# Developer section writer
def Developer_section_writer(Developer_section_dict, BSF_developer_transfers, Developer_tables, figure_count, table_count, dates_variables, paths_variables, DR):
    # Unpacking Dates variables
    figure_path = os.path.join(paths_variables['figure_path'], f'Figure{figure_count}.svg')
    
    cutoff = dates_variables['cutoff']
    last_month = dates_variables['last_month']
    #converts this month into a format which will work in the hyperlink
    hyperlink_month = dates_variables['hyperlink_month']

    Developer_remediation_table = Developer_tables['Developer_remediation_table']

    # Unpack variables
    Developer_cutoff = Developer_section_dict['Developer_cutoff']

    Developer_started_c_pct = Developer_section_dict['Developer_started_c_pct']
    Developer_signoff_c_pct = Developer_section_dict['Developer_signoff_c_pct']

    Developer_data_change_line = Developer_section_dict['Developer_data_change_line']

    Developer_number_of = Developer_section_dict['Developer_number_of']

    Developer_responsibility_total = Developer_section_dict['Developer_responsibility_total']

    Developer_life_critical_total = Developer_section_dict['Developer_life_critical_total']
    Developer_life_critical_total_line = Developer_section_dict['Developer_life_critical_total_line']

    Developer_signoff_c_no = Developer_section_dict['Developer_signoff_c_no']
    Developer_signoff_line = Developer_section_dict['Developer_signoff_line']

    Developer_complete_c_no = Developer_section_dict['Developer_complete_c_no']
    Developer_complete_c_pct = Developer_section_dict['Developer_complete_c_pct']

    Developer_started_c_no = Developer_section_dict['Developer_started_c_no']
    Developer_started_line = Developer_section_dict['Developer_started_line']

    Developer_plans_c_no = Developer_section_dict['Developer_plans_c_no']
    Developer_plans_c_pct = Developer_section_dict['Developer_plans_c_pct']
    Developer_plans_line = Developer_section_dict['Developer_plans_line']

    Developer_no_plans_c_no = Developer_section_dict['Developer_no_plans_c_no']
    Developer_no_plans_c_pct = Developer_section_dict['Developer_no_plans_c_pct']
    Developer_no_plans_line = Developer_section_dict['Developer_no_plans_line']

    Developer_cost = Developer_section_dict['Developer_cost']
    Developer_cost_change = Developer_section_dict['Developer_cost_change']

    Developer_dwellings_total = Developer_section_dict['Developer_dwellings_total']
    Developer_dwellings_started_c_no = Developer_section_dict['Developer_dwellings_started_c_no']

    Developer_forecast_timeframe = Developer_section_dict['Developer_forecast_timeframe']
    Developer_forecast_starts = Developer_section_dict['Developer_forecast_starts']
    Developer_forecast_completes = Developer_section_dict['Developer_forecast_completes']

    Developer_BSF_transfer_total = Developer_section_dict['Developer_BSF_transfer_total']
    Developer_BSF_transfer_total_line = Developer_section_dict['Developer_BSF_transfer_total_line']

    Developer_BSF_transfer_complete_c_no = Developer_section_dict['Developer_BSF_transfer_complete_c_no']
    Developer_BSF_transfer_complete_c_pct = Developer_section_dict['Developer_BSF_transfer_complete_c_pct']

    Developer_BSF_transfer_signoff_c_no = Developer_section_dict['Developer_BSF_transfer_signoff_c_no']
    Developer_BSF_transfer_signoff_c_pct = Developer_section_dict['Developer_BSF_transfer_signoff_c_pct']
    Developer_BSF_transfer_signoff_line = Developer_section_dict['Developer_BSF_transfer_signoff_line']

    Developer_BSF_transfer_started_c_no = Developer_section_dict['Developer_BSF_transfer_started_c_no']
    Developer_BSF_transfer_started_c_pct = Developer_section_dict['Developer_BSF_transfer_started_c_pct']
    Developer_BSF_transfer_started_line = Developer_section_dict['Developer_BSF_transfer_started_line']

    Developer_BSF_transfer_plans_nc_no = Developer_section_dict['Developer_BSF_transfer_plans_nc_no']
    Developer_BSF_transfer_plans_nc_pct = Developer_section_dict['Developer_BSF_transfer_plans_nc_pct']
    Developer_BSF_transfer_plans_line = Developer_section_dict['Developer_BSF_transfer_plans_line']

    Developer_BSF_transfer_no_plans_nc_no = Developer_section_dict['Developer_BSF_transfer_no_plans_nc_no']
    Developer_BSF_transfer_no_plans_nc_pct = Developer_section_dict['Developer_BSF_transfer_no_plans_nc_pct']
    Developer_BSF_transfer_no_plans_line = Developer_section_dict['Developer_BSF_transfer_no_plans_line']

    Developer_BSF_transfer_not_life_critical_nc_no = Developer_section_dict['Developer_BSF_transfer_not_life_critical_nc_no']
    Developer_BSF_transfer_not_life_critical_nc_pct = Developer_section_dict['Developer_BSF_transfer_not_life_critical_nc_pct']
    Developer_BSF_transfer_not_life_critical_line = Developer_section_dict['Developer_BSF_transfer_not_life_critical_line']

    Developer_11_18m_started_c_pct = Developer_section_dict['Developer_11_18m_started_c_pct']
    Developer_18m_started_c_pct = Developer_section_dict['Developer_18m_started_c_pct']

    Developer_cladding_defects_total = Developer_section_dict['Developer_cladding_defects_total']
    Developer_cladding_defects_total_line = Developer_section_dict['Developer_cladding_defects_total_line']

    Developer_cladding_defects_signoff_c_no = Developer_section_dict['Developer_cladding_defects_signoff_c_no']
    Developer_cladding_defects_signoff_c_pct = Developer_section_dict['Developer_cladding_defects_signoff_c_pct']
    Developer_cladding_defects_signoff_line = Developer_section_dict['Developer_cladding_defects_signoff_line']

    Developer_cladding_defects_complete_c_no = Developer_section_dict['Developer_cladding_defects_complete_c_no']
    Developer_cladding_defects_complete_c_pct = Developer_section_dict['Developer_cladding_defects_complete_c_pct']

    Developer_cladding_defects_started_c_no = Developer_section_dict['Developer_cladding_defects_started_c_no']
    Developer_cladding_defects_started_c_pct = Developer_section_dict['Developer_cladding_defects_started_c_pct']
    Developer_cladding_defects_started_line = Developer_section_dict['Developer_cladding_defects_started_line']

    Developer_cladding_defects_plans_c_no = Developer_section_dict['Developer_cladding_defects_plans_c_no']
    Developer_cladding_defects_plans_c_pct = Developer_section_dict['Developer_cladding_defects_plans_c_pct']
    Developer_cladding_defects_plans_line = Developer_section_dict['Developer_cladding_defects_plans_line']

    Developer_cladding_defects_no_plans_c_no = Developer_section_dict['Developer_cladding_defects_no_plans_c_no']
    Developer_cladding_defects_no_plans_c_pct = Developer_section_dict['Developer_cladding_defects_no_plans_c_pct']
    Developer_cladding_defects_no_plans_line = Developer_section_dict['Developer_cladding_defects_no_plans_line']

    Developer_self_reported_total = Developer_section_dict['Developer_self_reported_total']
    Developer_self_reported_total_line = Developer_section_dict['Developer_self_reported_total_line']

    Developer_self_reported_signoff_c_no = Developer_section_dict['Developer_self_reported_signoff_c_no']
    Developer_self_reported_signoff_c_pct = Developer_section_dict['Developer_self_reported_signoff_c_pct']

    Developer_self_reported_complete_c_no = Developer_section_dict['Developer_self_reported_complete_c_no']
    Developer_self_reported_complete_c_pct = Developer_section_dict['Developer_self_reported_complete_c_pct']

    Developer_self_reported_signoff_line = Developer_section_dict['Developer_self_reported_signoff_line']

    Developer_self_reported_started_c_no = Developer_section_dict['Developer_self_reported_started_c_no']
    Developer_self_reported_started_c_pct = Developer_section_dict['Developer_self_reported_started_c_pct']

    Developer_self_reported_plans_c_no = Developer_section_dict['Developer_self_reported_plans_c_no']
    Developer_self_reported_plans_c_pct = Developer_section_dict['Developer_self_reported_plans_c_pct']
    Developer_self_reported_plans_line = Developer_section_dict['Developer_self_reported_plans_line']

    Developer_self_reported_no_plans_c_no = Developer_section_dict['Developer_self_reported_no_plans_c_no']
    Developer_self_reported_no_plans_c_pct = Developer_section_dict['Developer_self_reported_no_plans_c_pct']
    Developer_self_reported_no_plans_line = Developer_section_dict['Developer_self_reported_no_plans_line']

    # Section Title
    paragraph = DR.add_paragraph('Developer-Led Remediation', style = 'Heading 2')

    # Paragraph
    text = f'Information in this section received by developers is correct as at {Developer_cutoff}{Developer_data_change_line}.'
    text +f'Information in this section received from other programmes that relate to developer-led remediation (where cladding remediation is being carried out in a government funded remediation programme and the developer will subsequently pay for the works) is correct as at {cutoff}.'
    DR.add_paragraph(text, style = 'Normal')

    # Paragraph
    text = 'The estimates in this section include some buildings which are also included in other sections of this data release e.g., those reported under the following sections: ‘ACM Remediation’, ‘Building Safety Fund’, ‘Cladding Safety Scheme’ and ‘Social Housing Sector’.'
    DR.add_paragraph(text, style = 'Normal')

    # Figure Title
    paragraph = DR.add_paragraph(style = 'Normal')
    text = f'Figure {figure_count}: {Developer_started_c_pct} of buildings in the developer remediation contract have either started or completed remediation works on life-critical fire safety risks, with {Developer_signoff_c_pct} having completed remediation works.'
    run = paragraph.add_run(text)
    run.bold = True
    
    # Figure
    DR.add_picture(os.path.join(paths_variables['figure_path'], f'Figure{figure_count}.svg'), width=Cm(17))
    figure_count += 1

    # Table caption
    text = f'Table {table_count}: Remediation status of buildings requiring works under the developer remediation contract, {cutoff}. Percentages may not sum to 100% due to rounding.'
    paragraph = DR.add_paragraph(style = 'Normal')
    run = paragraph.add_run(text)
    run.bold = True
    table_count += 1

    # Table
    table_data = Developer_remediation_table
    table_widths = [Cm(6.5), Cm(2.65), Cm(2.65), Cm(2.75), Cm(2.75)]
    table_heights = [Cm(1.12), Cm(0.55), Cm(1.13), Cm(0.55), Cm(1.13), Cm(1.13), Cm(0.55)]
    create_table(DR, table_data, table_widths, table_heights)

    # Heading
    paragraph = DR.add_paragraph('Developer remediation: key statistics', style = 'Heading 3')

    # Paragraph
    text = f'he estimates in this section are based on a combination of self-reported data submitted by developers and information that has been imputed from Building Safety Fund (BSF), Cladding Safety Scheme (CSS) and ACM programme data. A building is identified with life-critical fire safety defects if:'
    DR.add_paragraph(text, style = 'Normal')

    # Bullet point
    text = f'The developer has self-reported that works are (or were) required on the building'
    DR.add_paragraph(text, style = 'List Bullet')

    # Bullet point
    text = f'The developer has reported a remediation status of planned, started or complete'
    DR.add_paragraph(text, style = 'List Bullet')

    # Bullet point
    text = f'The building is eligible for funding in the BSF, Cladding Safety Scheme (CSS) or is being monitored under the ACM programme (including buildings developers have agreed to reimburse taxpayers for); '
    DR.add_paragraph(text, style = 'List Bullet')

    # Bullet point
    text = f'The building has had money paid out by the BSF or ACM remediation funds.'
    DR.add_paragraph(text, style = 'List Bullet')

    # Paragraph
    text = f'Furthermore, where a building has been identified with life-critical fire safety defects but the developer has not reported a defect type in the question ‘What do the issues relate to’ it is assumed that the defect type relates to cladding.'
    DR.add_paragraph(text, style = 'Normal')

    # Paragraph
    text = f'Other data, unless otherwise stated, is based on self-reported data by developers in their latest data report.'
    DR.add_paragraph(text, style = 'Normal')

    # Paragraph
    text = f'As at {cutoff}, {Developer_number_of} developers have signed the developer remediation contract. There are {Developer_responsibility_total} buildings covered by the developer remediation contract.'
    text += f'Of these, {Developer_life_critical_total} have been identified as having life-critical fire safety defects that developers are obligated to remediate or pay to remediate – {Developer_life_critical_total_line} since the {last_month} data release.'
    DR.add_paragraph(text, style = 'Normal')

    # Paragraph
    text = f'Of the {Developer_life_critical_total} buildings identified as having life-critical fire safety risks:'
    DR.add_paragraph(text, style = 'Normal')

    # Bullet point
    text = f'{Developer_signoff_c_no} ({Developer_signoff_c_pct}) are reported to have completed remediation – {Developer_signoff_line} since the {last_month} data release.'
    text += f'Of the {Developer_signoff_c_no} buildings that are reported to have completed remediation, {Developer_complete_c_no} buildings ({Developer_complete_c_pct} of all buildings with defects) are reported to have received building control sign-off.'
    DR.add_paragraph(text, style = 'List Bullet')

    # Bullet point
    text = f'{Developer_started_c_no} ({Developer_started_c_pct}) are reported to have started or completed remediation – {Developer_started_line} since the {last_month} data release.'
    DR.add_paragraph(text, style = 'List Bullet')
    # Bullet point
    text = f'{Developer_plans_c_no} ({Developer_plans_c_pct}) are reported to have not started remediation but have plans in place – {Developer_plans_line} since the {last_month} data release.'
    DR.add_paragraph(text, style = 'List Bullet')

    # Bullet point
    text = f'{Developer_no_plans_c_no} buildings ({Developer_no_plans_c_pct}) have not started and have no plans in place - {Developer_no_plans_line} since the {last_month} data release.'
    DR.add_paragraph(text, style = 'List Bullet')

    # Paragraph
    DR.add_paragraph('Although information from developers is received quarterly, these statistics are subject to change and are updated monthly as information from other programmes which relate to remediation is updated monthly.', style = 'Normal')

    # Paragraph 
    paragraph = DR.add_paragraph(style = 'Normal')
    text = f'The {Developer_life_critical_total} buildings identified as requiring remediation have an estimated cost of remediation of around £{Developer_cost} billion.'
    run = paragraph.add_run(text)
    if Developer_cost_change == 0:
        text = f' This has not changed compared to the {last_month} data release. This cost is based on the self-reported cost of works in the developer data report. However, if the cost is not known it is imputed based on the average known reported cost by developers by height band of building.'
        run = paragraph.add_run(text)
    else:
        text = f' This is an increase of £{Developer_cost_change}bn since reported in the {last_month} data release. This cost is based on the self-reported cost of works in the developer data report. However, if the cost is not known it is imputed based on the average known reported cost by developers by height band of building.'
        run = paragraph.add_run(text)
        run = paragraph.add_run('[ADD REASON FOR CHANGE]')
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # Paragraph
    text = f'There are an estimated {Developer_dwellings_total} dwellings in buildings with defects that developers are committed to remediate. Of these, there are an estimated {Developer_dwellings_started_c_no} dwellings in buildings that are reported as having either started or completed remediation works.'
    DR.add_paragraph(text, style = 'Normal')

    # Paragraph
    text = f'Based on start and completion dates reported by developers, {Developer_forecast_starts} buildings which have not yet started are reportedly expected to start works between {Developer_forecast_timeframe}, and {Developer_forecast_completes} buildings which have not yet completed are reportedly expected to complete their remediation between {Developer_forecast_timeframe}.'
    DR.add_paragraph(text, style = 'Normal')

    # Paragraph
    text = f'{Developer_BSF_transfer_total} buildings have transferred from the Building Safety Fund to developers, {Developer_BSF_transfer_total_line} since the {last_month} data release:'
    DR.add_paragraph(text, style = 'Normal')

    # Bullet point
    text = f'{Developer_BSF_transfer_signoff_c_no} ({Developer_BSF_transfer_signoff_c_pct}) have completed remediation - {Developer_BSF_transfer_signoff_line} since the {last_month} data release.'
    text+= f'Of these {Developer_BSF_transfer_signoff_c_no}, {Developer_BSF_transfer_complete_c_no} buildings ({Developer_BSF_transfer_complete_c_pct} of all buildings) are reported to have received building control sign-off.'
    DR.add_paragraph(text, style = 'List Bullet')

    # Bullet point
    text = f'{Developer_BSF_transfer_started_c_no} ({Developer_BSF_transfer_started_c_pct}) have started or completed remediation – {Developer_BSF_transfer_started_line} since the {last_month} data release.'
    DR.add_paragraph(text, style = 'List Bullet')

    # Bullet point
    text = f'{Developer_BSF_transfer_plans_nc_no} ({Developer_BSF_transfer_plans_nc_pct}) are reported to have not started remediation but have plans in place – {Developer_BSF_transfer_plans_line} since the {last_month} data release.'
    DR.add_paragraph(text, style = 'List Bullet')

    # Bullet point
    text = f'{Developer_BSF_transfer_no_plans_nc_no} ({Developer_BSF_transfer_no_plans_nc_pct}) are reported to have not started remediation and have no plans in place – {Developer_BSF_transfer_no_plans_line} since the {last_month} data release.'
    DR.add_paragraph(text, style = 'List Bullet')

    # Bullet point
    text = f'{Developer_BSF_transfer_not_life_critical_nc_no} ({Developer_BSF_transfer_not_life_critical_nc_pct}) have not been identified by developers as having life-critical fire-safety defects – {Developer_BSF_transfer_not_life_critical_line} since the {last_month} data release.'
    DR.add_paragraph(text, style = 'List Bullet')

    # Paragraph
    text = f'The {Developer_BSF_transfer_total} buildings in the developer remediation contract differs from the {BSF_developer_transfers} reported in the Building Safety Fund section of the data release, due to developers defining buildings differently to in the Building Safety Fund.'
    text = 'The same building structures are included in both sections of the release.'
    DR.add_paragraph(text, style = 'Normal')

    # Heading
    paragraph = DR.add_paragraph('Height breakdown', style = 'Heading 3')

    # Figure Title
    paragraph = DR.add_paragraph(style = 'Normal')
    text = f'Figure {figure_count}: {Developer_18m_started_c_pct} of the 18m+ buildings have started or completed remediation, compared to {Developer_11_18m_started_c_pct} of the 11-18m buildings.'
    run = paragraph.add_run(text)
    run.bold = True
    

    # Figure
    DR.add_picture(os.path.join(paths_variables['figure_path'], f'Figure{figure_count}.svg'), width=Cm(17))
    figure_count += 1
    
    # Heading
    paragraph = DR.add_paragraph('Cladding defects', style = 'Heading 3')

    # Paragraph
    text = f'The estimates in this section are based on a combination of self-reported data submitted by developers and information that has been imputed from Building Safety Fund (BSF) and ACM programme data (please see above for further details).'
    DR.add_paragraph(text, style = 'Normal')

    # Paragraph
    text = f'Where a building has been identified with life-critical fire safety defects but the developer has not reported a defect type in the question ‘What do the issues relate to’ it is assumed that the defect type relates to cladding.'
    DR.add_paragraph(text, style = 'Normal')

    # Paragraph
    text = f'Some remediation being undertaken by developers on buildings with life-critical fire safety risks relate to non-cladding defects.' 
    text += f'When excluding buildings reported to have only non-cladding defects, there are {Developer_cladding_defects_total} buildings which developers have reported as having unsafe cladding – {Developer_cladding_defects_total_line} since the {last_month} data release. Of these:'
    DR.add_paragraph(text, style = 'Normal')

    # Bullet point
    text = f'{Developer_cladding_defects_signoff_c_no} ({Developer_cladding_defects_signoff_c_pct}) are reported to have completed remediation– {Developer_cladding_defects_signoff_line} since the {last_month} data release.'
    text += f'Of the {Developer_cladding_defects_signoff_c_no} buildings that are reported to have completed remediation, {Developer_cladding_defects_complete_c_no} buildings ({Developer_cladding_defects_complete_c_pct} of all buildings with defects) are reported to have received building control sign-off.'
    DR.add_paragraph(text, style = 'List Bullet')

    # Bullet point
    text = f'{Developer_cladding_defects_started_c_no} ({Developer_cladding_defects_started_c_pct}) are reported to have started or completed remediation – {Developer_cladding_defects_started_line} since the {last_month} data release.'
    DR.add_paragraph(text, style = 'List Bullet')

    # Bullet point
    text = f'{Developer_cladding_defects_plans_c_no} ({Developer_cladding_defects_plans_c_pct}) are reported to have not started remediation but have plans in place – {Developer_cladding_defects_plans_line} since the {last_month} data release.'
    DR.add_paragraph(text, style = 'List Bullet')

    # Bullet point
    text = f'{Developer_cladding_defects_no_plans_c_no} ({Developer_cladding_defects_no_plans_c_pct}) are reported to have not started remediation but have plans in place – {Developer_cladding_defects_no_plans_line} since the {last_month} data release.'
    DR.add_paragraph(text, style = 'List Bullet')

    # Paragraph
    DR.add_paragraph('These estimates are subject to change each month due to the monitoring of buildings under the developer remediation contract in other remediation programmes.', style = 'Normal')

    # Heading
    paragraph = DR.add_paragraph('Self-reported information', style = 'Heading 3')

    # Paragraph 
    DR.add_paragraph('The figures reported for developers above combine information received directly from developers as well as information held by the department from other programmes which relate to remediation. \
                    This is to help better estimate a figure for buildings requiring remediation under the developer remediation contract. The figures reported above will also include buildings which are being remediated solely under a government remediation scheme, for which the developer will reimburse taxpayers.', style = 'Normal')

    # Paragraph
    text = f'The estimates in this section report on buildings which have been self-reported by developers as requiring remediation to life-critical fire safety risks. Furthermore, this section reports on buildings which are being remediated by the developer directly, rather than being remediated through a government fund and being reimbursed to taxpayers.'
    text += f'However, some buildings being remediated under a government remediation programme will be included if other relevant fire safety defects have been found, which were not eligible for a government remediation programme and which the developer is remediating themselves. Developers have self-reported that {Developer_self_reported_total} buildings require life-critical fire safety remediation, which will be directly remediated by the developer, {Developer_self_reported_total_line} since the {last_month} data release.'
    DR.add_paragraph(text, style = 'Normal')

    # Paragraph
    text = f'Buildings that are being remediated in the BSF, ACM or CSS programme, but are also reported to have other non-EWS life-critical fire safety defects by the developer will be included in these statistics. Of these {Developer_self_reported_total} buildings being directly remediated by the developer:'
    DR.add_paragraph(text, style = 'Normal')

    # Bullet point
    text = f'{Developer_self_reported_signoff_c_no} buildings ({Developer_self_reported_signoff_c_pct}) have completed remediation, {Developer_self_reported_signoff_line} since the {last_month} data release.'
    text += f'Of these, {Developer_self_reported_complete_c_no} buildings ({Developer_self_reported_complete_c_pct} of all buildings with defects) are reported to have received building control sign-off.'
    DR.add_paragraph(text, style = 'List Bullet')

    # Bullet point
    text = f'{Developer_self_reported_started_c_no} buildings ({Developer_self_reported_started_c_pct}) have started or completed remediation, {Developer_BSF_transfer_started_line} since the {last_month} data release.'
    DR.add_paragraph(text, style = 'List Bullet')

    # Bullet point
    text = f'{Developer_self_reported_plans_c_no} buildings ({Developer_self_reported_plans_c_pct}) have not started remediation but have a plan in place, {Developer_self_reported_plans_line} since the {last_month} data release.'
    DR.add_paragraph(text, style = 'List Bullet')

    # Bullet point
    text = f'{Developer_self_reported_no_plans_c_no} buildings ({Developer_self_reported_no_plans_c_pct}) have not started and have no plans in place, {Developer_self_reported_no_plans_line} since the {last_month} data release.'
    DR.add_paragraph(text, style = 'List Bullet')

    # Paragraph
    text = f'Further information on the progress developers have made regarding the buildings they’ve reported on is available in the accompanying '
    paragraph = DR.add_paragraph(text, style = 'Normal')
    add_hyperlink(paragraph, 'management information tables', f'https://www.gov.uk/government/publications/building-safety-remediation-monthly-data-release-{hyperlink_month}')
    paragraph.add_run(' and quarterly Developer remediation contract data release. From August 2025 these include an update against the commitments set out in the ')
    add_hyperlink(paragraph, 'joint plan to accelerate developer-led remediation and improve resident experience', f'https://www.gov.uk/government/publications/joint-plan-to-accelerate-developer-led-remediation-and-improve-resident-experience')
    paragraph.add_run(', which was published alongside the department Remediation Acceleration Plan (RAP).')

    #Paragraph
    text = 'Additionally, alongside the Building Safety Remediation Data Release, MHCLG publishes a ‘developer progress chart’ which allows you to compare the progress developers have made on determining whether works are required on buildings they are responsible for, as well as progress being made on buildings requiring works that have started on site.'
    text += 'This chart represents the self-reported information shown above and is published in the [accompanying dashboard](DASHBOARD LINK).'
    DR.add_paragraph(text, style = 'Normal')



    return figure_count, table_count