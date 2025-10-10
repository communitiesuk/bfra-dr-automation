# Developer_writer_headline.py
"""
Created on Wednesday 19 February 2025, 14:39:57

Author: Harry Simmons
"""

from docx import Document
import sys
import os

# Add the Utility folder to sys.path
folder_path = os.path.abspath(os.path.join(os.getcwd(), '..', 'Utility'))  # Replace 'folder_name' with the folder's name
sys.path.append(folder_path)

def Developer_headline_writer(Developer_headline_dict, dates_variables, DR):
    # Unpacking dates variables 
    cutoff = dates_variables['cutoff']
    last_month = dates_variables['last_month']

    # Unpacking  variables 
    Developer_cutoff = Developer_headline_dict['Developer_cutoff']
    Developer_life_critical_total = Developer_headline_dict['Developer_life_critical_total']
    Developer_life_critical_total_line = Developer_headline_dict['Developer_life_critical_total_line']
    Developer_started_c_no = Developer_headline_dict['Developer_started_c_no']
    Developer_started_c_pct = Developer_headline_dict['Developer_started_c_pct']
    Developer_started_line = Developer_headline_dict['Developer_started_line']
    Developer_signoff_c_no = Developer_headline_dict['Developer_signoff_c_no']
    Developer_signoff_c_pct = Developer_headline_dict['Developer_signoff_c_pct']
    Developer_signoff_line = Developer_headline_dict['Developer_signoff_line']
    Developer_cladding_defects_total = Developer_headline_dict['Developer_cladding_defects_total']
    Developer_cladding_defects_total_line = Developer_headline_dict['Developer_cladding_defects_total_line']
    Developer_cladding_defects_started_c_no = Developer_headline_dict['Developer_cladding_defects_started_c_no']
    Developer_cladding_defects_started_c_pct = Developer_headline_dict['Developer_cladding_defects_started_c_pct']
    Developer_cladding_defects_started_line = Developer_headline_dict['Developer_cladding_defects_started_line']
    Developer_cladding_defects_signoff_c_no = Developer_headline_dict['Developer_cladding_defects_signoff_c_no']
    Developer_cladding_defects_signoff_c_pct = Developer_headline_dict['Developer_cladding_defects_signoff_c_pct']
    Developer_cladding_defects_signoff_line = Developer_headline_dict['Developer_cladding_defects_signoff_line']

    # Headline Title
    text = f'Developer remediation – data received from developers as at {Developer_cutoff}. Data received from other programmes relating to developer-led remediation is as at {cutoff}.'
    paragraph = DR.add_paragraph(text, style = 'Heading 3')

    # Paragraph
    text = f'As at {cutoff}, {Developer_life_critical_total} buildings 11 metres and over in height have been identified as having life-critical fire safety defects (including cladding and non-cladding defects) which developers have committed to remediate or pay to remediate (where the cladding remediation works are being carried out in a government funded remediation programme), {Developer_life_critical_total_line} since reported in the {last_month} data release.'
    DR.add_paragraph(text, style = 'Normal')

    # Paragraph
    text = f'Of these, developers reported that {Developer_started_c_no} ({Developer_started_c_pct}) have either started or completed remediation works, {Developer_started_line} since reported in the {last_month} data release. Of these, {Developer_signoff_c_no} ({Developer_signoff_c_pct} of buildings) are reported to have completed remediation works, {Developer_signoff_line} since reported in the {last_month} data release.'
    DR.add_paragraph(text, style = 'Normal')

    # Paragraph 
    text = f'When excluding buildings reported with only non-cladding defects, there are {Developer_cladding_defects_total} buildings which developers have reported as having unsafe cladding, {Developer_cladding_defects_total_line} since reported in the {last_month} data release. Of which {Developer_cladding_defects_started_c_no} ({Developer_cladding_defects_started_c_pct}) are reported to have started or completed remediation, {Developer_cladding_defects_started_line} since reported in the {last_month} data release, including {Developer_cladding_defects_signoff_c_no} ({Developer_cladding_defects_signoff_c_pct} of buildings) which are reported to have completed remediation, {Developer_cladding_defects_signoff_line} since reported in the {last_month} data release.'
    DR.add_paragraph(text, style = 'Normal')