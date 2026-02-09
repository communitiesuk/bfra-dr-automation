# Social_section_writer.py
"""
Created on Wednesday 01 October 2025, 15:55:04

Author: Matthew Bandura
"""
from docx.shared import Cm
import os

from Utility.functions import create_table, add_hyperlink


def Social_section_writer(Social_section_dict, Social_tables, figure_count, table_count, dates_variables, paths_variables, DR):
    
    cutoff = dates_variables['cutoff']
    last_month = dates_variables['last_month']
    last_month_year = dates_variables['last_month_year']

    Social_remediation_table = Social_tables['Social_remediation_table']

    Social_total_no = Social_section_dict['Social_total_no']
    Social_total_change = Social_section_dict['Social_total_change']

    Social_completed_no = Social_section_dict['Social_completed_no']
    Social_completed_pct = Social_section_dict['Social_completed_pct']
    Social_completed_change = Social_section_dict['Social_completed_change']

    Social_started_no = Social_section_dict['Social_started_no']
    Social_started_pct = Social_section_dict['Social_started_pct']
    Social_started_change = Social_section_dict['Social_started_change']

    Social_not_started_no = Social_section_dict['Social_not_started_no']
    Social_not_started_pct = Social_section_dict['Social_not_started_pct']
    Social_not_started_change = Social_section_dict['Social_not_started_change']

    Social_unknown_no = Social_section_dict['Social_unknown_no']
    Social_unknown_pct = Social_section_dict['Social_unknown_pct']

    Social_SF_18m_starts_pct = Social_section_dict['Social_SF_18m_starts_pct']
    Social_SF_11m_starts_pct = Social_section_dict['Social_SF_11m_starts_pct']

    # Section Heading
    DR.add_paragraph('Social housing self-funded remediation', style = 'Heading 2')

    # Subheading
    DR.add_paragraph('Data quality', style = 'Heading 3')
    
    # Paragraph 1
    text = 'Homes England are continuing to work with registered providers to review records of their 11m+ buildings which have or had unsafe cladding to the NRS. While this work is underway, fluctuations in the remediation progress figures for social buildings where registered providers are self-funding cladding remediation is expected.'
    text += ' Once this work is complete, we expect to be able to provide more accurate and frequent information on social housing remediation progress. Further information on the data collection of social self-funded buildings is available in the technical note.'
    DR.add_paragraph(text, style = 'Normal')

    # Figure Title
    paragraph = DR.add_paragraph(style = 'Normal')
    text = f'Figure {figure_count}: {Social_started_pct} of social buildings identified to have unsafe cladding have started or completed remediation works, with {Social_completed_pct} (of identified buildings) having completed remediation works, including those awaiting building control sign-off.'
    run = paragraph.add_run(text)
    run.bold = True
 
    # Heading
    DR.add_paragraph('Social Housing remediation: Key statistics', style = 'Heading 3')

    # Figure
    DR.add_picture(os.path.join(paths_variables['figure_path'], f'Figure{figure_count}.svg'), width=Cm(17))
    figure_count += 1

    # Table caption
    text = f'Table {table_count}: Remediation status of social buildings with unsafe cladding, {cutoff}.'
    paragraph = DR.add_paragraph(style = 'Normal')
    run = paragraph.add_run(text)
    run.bold = True
    table_count += 1

    # Table
    table_data = Social_remediation_table
    table_widths = [Cm(6.5), Cm(2.65), Cm(2.65), Cm(2.75), Cm(2.75)]
    table_heights = [Cm(1.12), Cm(0.55), Cm(1.13), Cm(0.55), Cm(1.13), Cm(0.55)]
    create_table(DR, table_data, table_widths, table_heights)

    # Paragraph
    text = f'As at {cutoff}, there are {Social_total_no} 11m+ social buildings identified with unsafe cladding which registered providers are self-funding the remediation of, a change of {Social_total_change} since the end of {last_month} {last_month_year}.'
    text += 'This includes social buildings that registered providers reported to be self-funding the cladding remediation of in the NRS, that were not identified in another remediation programme, and 19 social buildings monitored in the ACM programme where registered providers have self-funded ACM remediation.'
    DR.add_paragraph(text, style = 'Normal')

    # Paragraph
    text = f'Of these {Social_total_no} 11m+ social buildings where registered providers are self-funding remediation:'
    DR.add_paragraph(text, style = 'Normal')

    # Bullet point 1
    text = f'{Social_completed_no} ({Social_completed_pct}%) are reported to have completed remediation – {Social_completed_change} since the end of {last_month} {last_month_year}.'
    DR.add_paragraph(text, style = 'List Bullet')

    # Bullet point 2
    text = f'{Social_started_no} ({Social_started_pct}%) are reported to have started or completed remediation – {Social_started_change} since the end of {last_month} {last_month_year}.'
    DR.add_paragraph(text, style = 'List Bullet')

    # Bullet point 3
    text = f'{Social_not_started_no} ({Social_not_started_pct}%) are reported to have started or completed remediation – {Social_not_started_change} since the end of {last_month} {last_month_year}.'
    DR.add_paragraph(text, style = 'List Bullet')

    # Bullet point 4
    text = f'The remediation status of {Social_unknown_no} ({Social_unknown_pct}%) buildings is currently unknown as registered providers are yet to provide remediation dates on the NRS. We expect buildings with unknown remediation status to be confirmed in coming months as further data is collected.'
    DR.add_paragraph(text, style = 'List Bullet')
    
    # Subheading
    DR.add_paragraph('Height breakdown', style = 'Heading 3')
 
    # Figure Title
    paragraph = DR.add_paragraph(style = 'Normal')
    text = f'Figure {figure_count}: {Social_SF_11m_starts_pct} of the 11-18m social self-funded buildings are reported to have started or completed remediation compared to {Social_SF_18m_starts_pct} of the 18m+ social self-funded buildings. '
    run = paragraph.add_run(text)
    run.bold = True

    # Figure
    DR.add_picture(os.path.join(paths_variables['figure_path'], f'Figure{figure_count}.svg'), width=Cm(17))
    figure_count += 1

    return figure_count, table_count
