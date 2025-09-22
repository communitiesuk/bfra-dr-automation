# DR_start_infrastructure.py
"""
Created on Monday 03 February 2025, 14:13:42

Author: Harry Simmons
"""

import docx


def DR_start(dates_variables, DR):
    # Unpack date variables
    cutoff = dates_variables['cutoff']
    publishing_date_0 = dates_variables['publishing_date_0']
    publishing_date_1 = dates_variables['publishing_date_1']

    # Document Title
    text = f'Building Safety Remediation Data Release: {cutoff}'
    paragraph = DR.add_paragraph(text, style = 'Heading 1')


    # Publishing date
    text = f'Published {publishing_date_0}'
    DR.add_paragraph(text, style = 'Normal')

    # Next publication date
    paragraph = DR.add_paragraph(style = 'Normal')
    run = paragraph.add_run('Date of next publication:')
    run.bold = True
    text = f' 9.30am on Thursday {publishing_date_1}'
    run = paragraph.add_run(text)

    # Paragraph
    paragraph = DR.add_paragraph(style='Normal')
    run = paragraph.add_run('All figures in this release can also be found in an ')
    run = paragraph.add_run('interactive dashboard [INSERT LINK]')

    run = paragraph.add_run('.')


    # Section Title
    paragraph = DR.add_paragraph('Headlines', style = 'Heading 2')

DR = docx.Document()

dates_variables = {
    'cutoff': 'dummy',
    'publishing_date_0': 'dummy_0',
    'publishing_date_1': 'dummy_1'
}



DR_start(dates_variables, DR)
