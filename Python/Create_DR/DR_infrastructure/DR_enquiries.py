# DR_start_infrastructure.py
"""
Created on Thursday 13 March 2025, 10:05:24

Author: Harry Simmons
"""

import docx
from docx.enum.text import WD_COLOR_INDEX
import sys
import os

folder_path = os.path.abspath(os.path.join(os.getcwd(), '..', 'Utility'))  # Replace 'folder_name' with the folder's name
sys.path.append(folder_path)

def DR_enquiries(dates_variables, DR):
    # Unpack date variables
    publishing_date_1 = dates_variables['publishing_date_1']

    # Section Title 
    paragraph = DR.add_paragraph('Enquiries', style = 'Heading 2')

    # Contact info
    paragraph = DR.add_paragraph(style = 'Normal')
    run = paragraph.add_run('Contact:')
    run.bold = True
    run = paragraph.add_run(' BuildingSafetyData2@communities.gov.uk\n')
    run = paragraph.add_run('Media enquiries:')
    run.bold = True
    run = paragraph.add_run(' 0303 444 1209\nNewsDesk@communities.gov.uk')

    # Heading
    paragraph = DR.add_paragraph('User Engagement', style = 'Heading 3')

    # Paragraph
    DR.add_paragraph('We are committed to improving and broadening this data release further in the months ahead and would welcome feedback both on the revised content of this data release and what could be further done in the future. Please contact BuildingSafetyData2@communities.gov.uk', style = 'Normal')

    # Heading
    paragraph = DR.add_paragraph('Dates of future publications', style = 'Heading 3')

    # Paragraph 
    DR.add_paragraph('The publication dates for the next three months are:', style = 'Normal')

    # Bullet point
    DR.add_paragraph(publishing_date_1, style = 'List Bullet')

    # Bullet point
    paragraph = DR.add_paragraph(style = 'List Bullet')
    run = paragraph.add_run('[ADD DATE]')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    # Bullet point
    paragraph = DR.add_paragraph(style = 'List Bullet')
    run = paragraph.add_run('[ADD DATE]')
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW