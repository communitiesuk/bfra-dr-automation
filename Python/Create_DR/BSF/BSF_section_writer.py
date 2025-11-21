# BSF_writer_section.py
"""
Created on Thursday 09 January 2025, 10:39:20

Author: Harry Simmons
"""

from docx.shared import Cm
import sys
import os

# Add the Utility folder to sys.path
folder_path = os.path.abspath(os.path.join(os.getcwd(), '..', 'Utility'))  # Replace 'folder_name' with the folder's name
sys.path.append(folder_path)

from Utility.functions import create_table
import Utility.docx_svg_patch

def BSF_section_writer(BSF_section_dict, BSF_tables, figure_count, table_count, dates_variables, paths_variables, DR):
    
    # Unpacking date variables
    cutoff = dates_variables['cutoff']
    last_month = dates_variables['last_month']
    last_year_month = dates_variables['last_year_month']
    

    BSF_remediation_table = BSF_tables['BSF_remediation_table']

    # Unpacking regular variables
    BSF_signoff_c_pct = BSF_section_dict['BSF_signoff_c_pct']
    BSF_started_c_pct = BSF_section_dict['BSF_started_c_pct']

    BSF_BSF_5_total = BSF_section_dict['BSF_BSF_5_total']

    BSF_CSS_transfers_this_month = BSF_section_dict['BSF_CSS_transfers_this_month']
    BSF_CSS_transfers_line = BSF_section_dict['BSF_CSS_transfers_line']
    BSF_CSS_transfers_not_started = BSF_section_dict['BSF_CSS_transfers_not_started']

    BSF_CSS_retained = BSF_section_dict['BSF_CSS_retained']

    BSF_ineligible = BSF_section_dict['BSF_ineligible']
    BSF_withdrawn = BSF_section_dict['BSF_withdrawn']
    BSF_developer_transfers = BSF_section_dict['BSF_developer_transfers']
    BSF_insufficient_evidence = BSF_section_dict['BSF_insufficient_evidence']

    BSF_BSF_1_total = BSF_section_dict['BSF_BSF_1_total']

    BSF_developer_reimbursed_word = BSF_section_dict['BSF_developer_reimbursed_word']
    BSF_developer_anticipated_word = BSF_section_dict['BSF_developer_anticipated_word']
    BSF_developer_reimbursed_pct = BSF_section_dict['BSF_developer_reimbursed_pct']
    BSF_developer_anticipated_pct = BSF_section_dict['BSF_developer_anticipated_pct']

    BSF_FRAEW = BSF_section_dict['BSF_FRAEW']
    BSF_CAN = BSF_section_dict['BSF_CAN']

    BSF_started_c_no = BSF_section_dict['BSF_started_c_no']
    BSF_started_c_line = BSF_section_dict['BSF_started_c_line']

    BSF_started_nc_no = BSF_section_dict['BSF_started_nc_no']
    BSF_started_nc_pct = BSF_section_dict['BSF_started_nc_pct']

    BSF_signoff_c_no = BSF_section_dict['BSF_signoff_c_no']
    BSF_signoff_c_line = BSF_section_dict['BSF_signoff_c_line']

    BSF_complete_nc_no = BSF_section_dict['BSF_complete_nc_no']
    BSF_complete_nc_pct = BSF_section_dict['BSF_complete_nc_pct']

    BSF_not_yet_started = BSF_section_dict['BSF_not_yet_started']

    BSF_plans_nc_no = BSF_section_dict['BSF_plans_nc_no']
    BSF_plans_nc_pct = BSF_section_dict['BSF_plans_nc_pct']

    BSF_intent_nc_no = BSF_section_dict['BSF_intent_nc_no']
    BSF_intent_nc_pct = BSF_section_dict['BSF_intent_nc_pct']

    BSF_dwellings = BSF_section_dict['BSF_dwellings']

    BSF_fig_eligible_change = BSF_section_dict['BSF_fig_eligible_change']
    BSF_fig_started_c_change = BSF_section_dict['BSF_fig_started_c_change']
    BSF_fig_completed_c_change = BSF_section_dict['BSF_fig_completed_c_change']

    BSF_social_started_pct = BSF_section_dict['BSF_social_started_pct']
    BSF_private_started_pct = BSF_section_dict['BSF_private_started_pct']
    BSF_social_complete_pct = BSF_section_dict['BSF_social_complete_pct']
    BSF_private_complete_pct = BSF_section_dict['BSF_private_complete_pct']

    # Section Title 
    paragraph = DR.add_paragraph('Building Safety Fund', style = 'Heading 2')

    # Paragraph
    text = f'Information in this section is correct as at {cutoff} and shows a monthly update from the previous publication.'
    paragraph = DR.add_paragraph(text, style = 'Normal')

    # Figure Title
    text = f'Figure {figure_count}: {BSF_started_c_pct} of buildings proceeding with an application for funding in the BSF have started or completed remediation, with {BSF_signoff_c_pct} having completed remediation.'
    paragraph = DR.add_paragraph(style = 'Normal')
    run = paragraph.add_run(text)
    run.bold = True
    

    # Figure 
    DR.add_picture(os.path.join(paths_variables['figure_path'], f'Figure{figure_count}.svg'), width=Cm(17))
    figure_count += 1

    # Table caption
    text = f'Table {table_count}: Remediation status of buildings within the Building Safety Fund, {cutoff}'
    paragraph = DR.add_paragraph(style = 'Normal')
    run = paragraph.add_run(text)
    run.bold = True
    table_count += 1

    # Table
    table_data = BSF_remediation_table
    table_widths = [Cm(6.5), Cm(2.65), Cm(2.65), Cm(2.75), Cm(2.75)]
    table_heights = [Cm(1.12), Cm(0.55), Cm(1.13), Cm(0.55), Cm(0.55), Cm(0.55), Cm(0.55), Cm(0.55), Cm(0.55)]
    create_table(DR, table_data, table_widths, table_heights)

    # Heading
    paragraph = DR.add_paragraph('BSF remediation: key statistics', style = 'Heading 3')

    # Paragraph
    text = f'As at {cutoff}, {BSF_BSF_5_total} buildings were assessed as eligible and are proceeding with an application for funding from the Building Safety Fund. The remaining {BSF_BSF_1_total} buildings registered with the fund are either ineligible ({BSF_ineligible}), withdrawn ({BSF_withdrawn}), in review or have given insufficient evidence ({BSF_insufficient_evidence}), have transferred to the Cladding Safety Scheme ({BSF_CSS_transfers_this_month}), or have been retained by the Cladding Safety Scheme ({BSF_CSS_retained}).'
    paragraph = DR.add_paragraph(text, style = 'Normal')

    # Paragraph
    text = f'Of the {BSF_BSF_1_total} buildings that registered with the Building Safety Fund and are not currently proceeding with an application for funding, {BSF_developer_transfers} buildings which were assessed as eligible have been transferred to developers.'
    paragraph = DR.add_paragraph(text, style = 'Normal')

    # Paragraph
    text = f'As at {cutoff}, {BSF_CSS_transfers_this_month} buildings have been transferred to the Cladding Safety Scheme for their remediation and are progressing through the CSS fund, {BSF_CSS_transfers_line} buildings since reported at the end of {last_month}. Of these, {BSF_CSS_transfers_not_started} buildings had not started remediation works before transferring to the CSS.'
    paragraph = DR.add_paragraph(text, style = 'Normal')

    # Paragraph
    text = f'The number of buildings reported as transferred from the BSF to the CSS in this section of the data release may not be the same as the number of eligible CSS buildings that have transferred from the BSF in the CSS section of the data release. This is because different programmes may define buildings differently due to how they operate. One building has also withdrawn from the CSS since transferring from the BSF.'
    paragraph = DR.add_paragraph(text, style = 'Normal')

    # Paragraph
    text = f'Of the {BSF_BSF_5_total} buildings proceeding with an application for funding in the Building Safety Fund:'
    paragraph = DR.add_paragraph(text, style = 'Normal')

    # Bullet point
    text = f'{BSF_developer_reimbursed_word.capitalize()} buildings ({BSF_developer_reimbursed_pct}) are remaining in the fund with developers set to reimburse the cost of remediation.'
    paragraph = DR.add_paragraph(text, style = 'List Bullet')

    # Bullet point
    text = f'{BSF_developer_anticipated_word.capitalize()} buildings ({BSF_developer_anticipated_pct}) are anticipated to be transferred to developers.'
    paragraph = DR.add_paragraph(text, style = 'List Bullet')

    # Paragraph
    paragraph = DR.add_paragraph(style = 'Normal')
    text = f'Of the {BSF_BSF_5_total} buildings proceeding with an application for funding, {BSF_FRAEW} buildings have been assessed with a Fire Risk Appraisal for External Walls (FRAEW), and {BSF_CAN} have been assessed under the BSF 2020 CAN criteria. It is possible for a building to submit a new application to the fund after having completed remediation, if it originally applied under the BSF 2020 CAN criteria. Therefore, buildings could be double counted. As of {cutoff} one building had undertaken remediation work from both application periods, so is double counted. Further details are available in the '
    run = paragraph.add_run(text)
    run = paragraph.add_run('technical note')
    run = paragraph.add_run('.')

    # Paragraph
    text = f'Of the {BSF_BSF_5_total} buildings proceeding with an application for funding, {BSF_started_c_no} buildings ({BSF_started_c_pct}) have either started or completed remediation works – {BSF_started_c_line} since the end of {last_month}.'
    paragraph = DR.add_paragraph(text)

    # Paragraph
    text = f'Of the {BSF_started_c_no} buildings that have started or completed remediation as at {cutoff}:'
    paragraph = DR.add_paragraph(text, style = 'Normal')

    # Bullet point
    text = f'{BSF_started_nc_no} buildings ({BSF_started_nc_pct} of all buildings) have started remediation.'
    paragraph = DR.add_paragraph(text, style = 'List Bullet')

    # Bullet point
    text = f'{BSF_signoff_c_no} buildings ({BSF_signoff_c_pct} of all buildings) have completed remediation – {BSF_signoff_c_line} since the end of {last_month}. Of these, {BSF_complete_nc_no} buildings ({BSF_complete_nc_pct} of all buildings) have received building control sign off.'
    paragraph = DR.add_paragraph(text, style = 'List Bullet')

    # Paragraph
    text = f'There are {BSF_not_yet_started} eligible buildings proceeding with an application for funding that have not started remediation, of which:'
    paragraph = DR.add_paragraph(text, style = 'Normal')

    # Bullet point
    text = f'{BSF_plans_nc_no} ({BSF_plans_nc_pct} of all buildings) have remediation plans in place.'
    paragraph = DR.add_paragraph(text, style = 'List Bullet')

    # Bullet point
    text = f'{BSF_intent_nc_no} ({BSF_intent_nc_pct} of all buildings) have reported an intent to remediate.'
    paragraph = DR.add_paragraph(text, style = 'List Bullet')

    # Paragraph
    text = f'There are an estimated {BSF_dwellings} dwellings in buildings that are eligible and proceeding with an application for funding in the BSF.'
    paragraph = DR.add_paragraph(text, style = 'Normal')

    # Heading
    paragraph = DR.add_paragraph('BSF remediation progress over time', style = 'Heading 3')

    # Figure Title
    text = f'Figure {figure_count}: {BSF_fig_started_c_change.capitalize()} buildings proceeding with an application for funding in the BSF have started or completed remediation since the end of {last_year_month}'
    paragraph = DR.add_paragraph(style = 'Normal')
    run = paragraph.add_run(text)
    run.bold = True
    

    # Figure
    DR.add_picture(os.path.join(paths_variables['figure_path'], f'Figure{figure_count}.svg'), width=Cm(17), height=Cm(15))
    figure_count += 1

    # Paragraph
    text = f'Since the end of {last_year_month}:'
    paragraph = DR.add_paragraph(text, style = 'Normal')

    # Bullet Point
    text = f'{BSF_fig_eligible_change.capitalize()} eligible buildings are proceeding with an application for funding from the Building Safety Fund.'
    paragraph = DR.add_paragraph(text, style = 'List Bullet')

    # Bullet Point
    text = f'{BSF_fig_started_c_change.capitalize()} eligible buildings have started or completed remediation, and {BSF_fig_completed_c_change} eligible buildings have completed remediation.'
    paragraph = DR.add_paragraph(text, style = 'List Bullet')

    # Heading
    paragraph = DR.add_paragraph('BSF remediation by tenure', style = 'Heading 3')

    # Figure  Title
    paragraph = DR.add_paragraph(style = 'Normal')
    if BSF_social_started_pct == BSF_private_started_pct:
        text = f'Figure {figure_count}: {BSF_social_started_pct} of private sector buildings and social sector buildings have started or completed remediation, with {BSF_private_complete_pct} of private sector buildings having completed remediation and {BSF_social_complete_pct} of social sector buildings.'
    else:
        text = f'Figure {figure_count}: {BSF_social_started_pct} of social sector buildings in the BSF, including buildings with a financial viability claim, have started or completed remediation compared to {BSF_private_started_pct} of private sector buildings.'
    run = paragraph.add_run(text)
    run.bold = True
    

    # Figure
    DR.add_picture(os.path.join(paths_variables['figure_path'], f'Figure{figure_count}.svg'), width=Cm(17), height=Cm(8))
    figure_count += 1
    
    return figure_count, table_count