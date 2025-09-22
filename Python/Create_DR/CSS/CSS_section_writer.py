# CSS_writer_section.py
"""
Created on Friday 17 January 2025, 14:32:06

Author: Harry Simmons
"""

import docx
from docx.shared import Cm
import sys
import os

# Add the Utility folder to sys.path
folder_path = os.path.abspath(os.path.join(os.getcwd(), '..', 'Utility'))  # Replace 'folder_name' with the folder's name
sys.path.append(folder_path)

from Utility.functions import create_table
import Utility.docx_svg_patch

def CSS_section_writer(CSS_section_dict, CSS_tables, figure_count, table_count, dates_variables, paths_variables, DR):
    # Unpacking date variables
    figure_path = os.path.join(paths_variables['figure_path'], f'Figure{figure_count}.svg')
    cutoff = dates_variables['cutoff']
    last_month = dates_variables['last_month']
    this_month = dates_variables['this_month']

    CSS_remediation_table = CSS_tables['CSS_remediation_table']

    CSS_total_total = CSS_section_dict['CSS_total_total']
    CSS_pre_eligible = CSS_section_dict['CSS_pre_eligible']
    CSS_eligible_total = CSS_section_dict['CSS_eligible_total']
    CSS_started_c_no = CSS_section_dict['CSS_started_c_no']
    CSS_started_c_pct = CSS_section_dict['CSS_started_c_pct']
    CSS_completed_c_no = CSS_section_dict['CSS_completed_c_no']
    CSS_completed_c_pct = CSS_section_dict['CSS_completed_c_pct']
    CSS_pre_eligible_total = CSS_section_dict['CSS_pre_eligible_total']
    CSS_pre_application = CSS_section_dict['CSS_pre_application']
    CSS_eligible_total_line = CSS_section_dict['CSS_eligible_total_line']
    CSS_BSF_transfer = CSS_section_dict['CSS_BSF_transfer']
    CSS_BSF_transfer_line = CSS_section_dict['CSS_BSF_transfer_line']
    CSS_GFA = CSS_section_dict['CSS_GFA']
    CSS_GFA_pct = CSS_section_dict['CSS_GFA_pct']
    CSS_GFA_line = CSS_section_dict['CSS_GFA_line']
    CSS_PTSP = CSS_section_dict['CSS_PTSP']
    CSS_PTSP_pct = CSS_section_dict['CSS_PTSP_pct']
    CSS_PTSP_line = CSS_section_dict['CSS_PTSP_line']
    CSS_started_c_line = CSS_section_dict['CSS_started_c_line']
    CSS_completed_c_line = CSS_section_dict['CSS_completed_c_line']
    CSS_northern_ireland = CSS_section_dict['CSS_northern_ireland']
    CSS_18m_c_pct = CSS_section_dict['CSS_18m_c_pct']
    CSS_11_18m_c_pct = CSS_section_dict['CSS_11_18m_c_pct']
    CSS_social_c_pct = CSS_section_dict['CSS_social_c_pct']
    CSS_private_c_pct = CSS_section_dict['CSS_private_c_pct']

    # Section Title 
    paragraph = DR.add_paragraph('Cladding Safety Scheme', style = 'Heading 2')

    # Paragraph
    text = f'Information in this section is correct as at {cutoff} and shows a monthly update from the previous publication.'
    paragraph = DR.add_paragraph(text, style = 'Normal')

    # Figure Title
    paragraph = DR.add_paragraph(style = 'Normal')
    text = f'Figure {figure_count}: There are {CSS_total_total} buildings at different stages of the Cladding Safety Scheme, including {CSS_pre_eligible} buildings with live applications and {CSS_eligible_total} eligible buildings, of which {CSS_started_c_no} buildings have started or completed works, with {CSS_completed_c_no} having completed works.'
    run = paragraph.add_run(text)
    run.bold = True
    
    # Figure
    DR.add_picture(figure_path, width=Cm(17))
    figure_count += 1

    # Table caption
    text = f'Table {table_count}: Remediation status for buildings within the Cladding Safety Scheme, {cutoff}'
    paragraph = DR.add_paragraph(style = 'Normal')
    run = paragraph.add_run(text)
    run.bold = True
    table_count += 1

    # Table
    table_data = CSS_remediation_table
    table_widths = [Cm(6.5), Cm(2.65), Cm(2.65), Cm(2.75), Cm(2.75)]
    table_heights = [Cm(1.12), Cm(0.55), Cm(0.55), Cm(0.55), Cm(0.55)]
    create_table(DR, table_data, table_widths, table_heights)

    # Heading
    paragraph = DR.add_paragraph('Cladding Safety Scheme: Key statistics', style = 'Heading 3')

    # Paragraph
    text = f'As at the end of {this_month}, there were {CSS_total_total} buildings in the different stages of the Cladding Safety Scheme ({CSS_eligible_total} eligible buildings and {CSS_pre_eligible_total} pre-eligible buildings). Of the {CSS_pre_eligible_total} pre-eligible buildings:'
    paragraph = DR.add_paragraph(text, style = 'Normal')

    # Bullet point
    text = f'{CSS_pre_application} buildings are in the pre-application stage.'
    paragraph = DR.add_paragraph(text, style = 'List Bullet')

    # Bullet point
    text = f'{CSS_pre_eligible} buildings have a live application and are progressing through the eligibility stages.'
    paragraph = DR.add_paragraph(text, style = 'List Bullet')

    # Paragraph
    text = f'There are {CSS_eligible_total} eligible buildings as at the end of {this_month}, {CSS_eligible_total_line} since the end of {last_month}. {CSS_BSF_transfer_line.capitalize()} of the newly eligible buildings this month have been transferred from the BSF. In total, {CSS_BSF_transfer} of the eligible buildings in the CSS have been transferred from the BSF.'
    paragraph = DR.add_paragraph(text, style = 'Normal')

    # Paragraph
    text = f'The number of eligible CSS buildings reported as transferred from the BSF to the CSS in this section of the data release may not be the same as the number of buildings that have transferred from the BSF to the CSS in the BSF section of the data release. This is because different programmes may define buildings differently due to how they operate. One building has also withdrawn from the CSS since transferring from the BSF.'
    paragraph = DR.add_paragraph(text, style = 'Normal')

    # Paragraph
    text = f'Of the {CSS_eligible_total} eligible buildings:'
    paragraph = DR.add_paragraph(text, style = 'Normal')

    # Bullet point
    text = f'{CSS_GFA} buildings ({CSS_GFA_pct} of eligible buildings) have signed Grant Funding Agreements (GFA) - {CSS_GFA_line} since the end of {last_month}.'
    paragraph = DR.add_paragraph(text, style = 'List Bullet')

    # Bullet point
    text = f'Of these, {CSS_PTSP} buildings ({CSS_PTSP_pct} of all eligible buildings) have received a pre-tender support payment – {CSS_PTSP_line} since the end of {last_month}.'
    paragraph = DR.add_paragraph(text, style = 'List Bullet')

    # Bullet point
    text = f'{CSS_started_c_no} buildings ({CSS_started_c_pct} of all eligible buildings) have started or completed remediation works on site – {CSS_started_c_line} since the end of {last_month}.'
    paragraph = DR.add_paragraph(text, style = 'List Bullet')

    # Bullet point
    text = f'Of these, {CSS_completed_c_no} buildings ({CSS_completed_c_pct} of all eligible buildings) have completed remediation works on site, including those awaiting building control sign-off – {CSS_completed_c_line} since the end of {last_month}.'
    paragraph = DR.add_paragraph(text, style = 'List Bullet')

    # Paragraph
    text = f'{CSS_northern_ireland.capitalize()} eligible buildings are in Northern Ireland, with the remainder in England.'
    paragraph = DR.add_paragraph(text, style = 'Normal')

    # Heading
    paragraph = DR.add_paragraph('Height breakdown', style = 'Heading 3')

    # Figure Title
    paragraph = DR.add_paragraph(style = 'Normal')
    text = f'Figure {figure_count}: {CSS_11_18m_c_pct} of eligible 11-18m buildings in the CSS have started or completed remediation, compared to {CSS_18m_c_pct} of eligible 18m+ buildings.'
    run = paragraph.add_run(text)
    run.bold = True

    # Figure
    DR.add_picture(figure_path, width=Cm(17))
    figure_count += 1

    # Heading
    paragraph = DR.add_paragraph('Tenure breakdown', style = 'Heading 3')

    # Figure Title
    paragraph = DR.add_paragraph(style = 'Normal')
    text = f'Figure {figure_count}: {CSS_social_c_pct} of eligible social sector buildings in the CSS have started or completed remediation, compared to {CSS_private_c_pct} of eligible private sector buildings.'
    run = paragraph.add_run(text)
    run.bold = True
   

    # Figure
    DR.add_picture(figure_path, width=Cm(17))
    figure_count += 1
    return figure_count, table_count