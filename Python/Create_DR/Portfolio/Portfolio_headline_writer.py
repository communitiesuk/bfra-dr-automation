# Portfolio_writer_headline.py
"""
Created on Monday 03 February 2025, 11:21:36

Author: Harry Simmons
"""

from docx.shared import Cm
import os
import Utility.docx_svg_patch

from Utility.functions import add_hyperlink

def Portfolio_headline_writer(Portfolio_headline_dict, Estimates_headline_dict, figure_count, dates_variables, paths_variables, DR):
    last_month = dates_variables['last_month']
    this_month = dates_variables['this_month']
    
    last_month_year = dates_variables['last_month_year']
    hyperlink_month = dates_variables['hyperlink_month']


    Portfolio_total = Portfolio_headline_dict['Portfolio_total']
    Portfolio_total_line = Portfolio_headline_dict['Portfolio_total_line']
    Portfolio_total_since_oct_23 = Portfolio_headline_dict['Portfolio_total_since_oct_23']
    Portfolio_started_c_no = Portfolio_headline_dict['Portfolio_started_c_no']
    Portfolio_started_c_pct = Portfolio_headline_dict['Portfolio_started_c_pct']
    Portfolio_completed_c_no = Portfolio_headline_dict['Portfolio_completed_c_no']
    Portfolio_completed_c_pct = Portfolio_headline_dict['Portfolio_completed_c_pct']
    Estimates_11m_remaining_low = Portfolio_headline_dict['Estimates_11m_remaining_low']
    Estimates_11m_remaining_high = Portfolio_headline_dict['Estimates_11m_remaining_high']

    Estimates_11m_proportion_of_low_estimate = Estimates_headline_dict['Estimates_11m_proportion_of_low_estimate']
    Estimates_11m_proportion_of_high_estimate = Estimates_headline_dict['Estimates_11m_proportion_of_high_estimate']


    # Headline Title
    text = f'Overall remediation'
    paragraph = DR.add_paragraph(text, style = 'Heading 3')

    # Paragraph
    text = f'As at the end of {this_month}, there are {Portfolio_total} residential buildings 11 metres and over in height identified with unsafe cladding whose remediation progression is being reported on in this release, {Portfolio_total_line} since the end of {last_month} {last_month_year}. This is an estimated {Estimates_11m_proportion_of_high_estimate}-{Estimates_11m_proportion_of_low_estimate} of all buildings 11 metres and over in height expected to be remediated as part of MHCLG’s remediation programmes.'
    DR.add_paragraph(text, style = 'Normal')

    # Paragraph
    text = f'Since the department first began reporting on all five remediation programmes in October 2023, {Portfolio_total_since_oct_23} buildings with unsafe cladding are being reported on in this release.'
    DR.add_paragraph(text, style = 'Normal')

    # Paragraph
    text = f'Overall, {Portfolio_started_c_no} buildings ({Portfolio_started_c_pct}) have either started or completed remediation works. Of these, {Portfolio_completed_c_no} buildings ({Portfolio_completed_c_pct}) have completed remediation works.'
    DR.add_paragraph(text, style = 'Normal')

    # Figure Title
    paragraph = DR.add_paragraph(style = 'Normal')
    text = f'Figure {figure_count}: {Portfolio_total} 11m+ buildings have been identified with unsafe cladding, and there are an estimated {Estimates_11m_remaining_low}-{Estimates_11m_remaining_high} 11m+ buildings expected to be remediated as part of MHCLG\'s remediation programmes yet to identify.'
    run = paragraph.add_run(text)
    run.bold = True
  
    # Figure
    DR.add_picture(os.path.join(paths_variables['figure_path'], f'Figure{figure_count}.svg'), width=Cm(17))
    figure_count += 1

    # Figure Title
    paragraph = DR.add_paragraph(style = 'Normal')
    text = f'Figure {figure_count}: Of the {Portfolio_total} buildings identified with unsafe cladding, {Portfolio_started_c_no} ({Portfolio_started_c_pct}) have started or completed remediation works, of which {Portfolio_completed_c_no} ({Portfolio_completed_c_pct}) have completed remediation works. This includes remediation progress on high rise (18m+) and mid-rise (11-18m) buildings in height.'
    run = paragraph.add_run(text)
    run.bold = True
  

    # Figure
    DR.add_picture(os.path.join(paths_variables['figure_path'], f'Figure{figure_count}.svg'), width=Cm(17))
    figure_count += 1

    # Figure caption
    paragraph = DR.add_paragraph('Note: The total number of buildings identified with unsafe cladding, reported in the ', style= 'Normal')
    add_hyperlink(paragraph, 'Overall Remediation section', f'https://www.gov.uk/government/publications/building-safety-remediation-monthly-data-release-{hyperlink_month}/building-safety-remediation-monthly-data-release-{hyperlink_month}#overall-remediation-progress')                 
    paragraph.add_run(' of the data release, does not sum to the total number of buildings in each remediation programme, reported in each respective section of the data release. This is due to some buildings appearing in more than one remediation programme.')

    return figure_count