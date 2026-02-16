# Portfolio_writer_section.py
"""
Created on Monday 27 January 2025, 09:15:26

Author: Harry Simmons
"""

from docx.shared import Cm
import os

from Utility.functions import create_table
import Utility.docx_svg_patch


def Portfolio_section_writer(Portfolio_section_dict, Portfolio_tables, figure_count, table_count, dates_variables, paths_variables, DR):
    # Unpacking date variables
    placeholder_path = paths_variables['placeholder_path']
    
    cutoff = dates_variables['cutoff']
    last_month = dates_variables['last_month']
    last_year_month = dates_variables['last_year_month']
    last_month_year = dates_variables['last_month_year']
    
    Portfolio_remediation_table = Portfolio_tables['Portfolio_remediation_table']

    Portfolio_started_c_no = Portfolio_section_dict['Portfolio_started_c_no']
    Portfolio_started_c_pct = Portfolio_section_dict['Portfolio_started_c_pct']
    Portfolio_completed_c_no = Portfolio_section_dict['Portfolio_completed_c_no']
    Portfolio_completed_c_pct = Portfolio_section_dict['Portfolio_completed_c_pct']
    Portfolio_total = Portfolio_section_dict['Portfolio_total']
    Portfolio_started_nc_no = Portfolio_section_dict['Portfolio_started_nc_no']
    Portfolio_started_nc_pct = Portfolio_section_dict['Portfolio_started_nc_pct']
    Portfolio_in_programme_nc_no = Portfolio_section_dict['Portfolio_in_programme_nc_no']
    Portfolio_in_programme_nc_pct = Portfolio_section_dict['Portfolio_in_programme_nc_pct']
    Portfolio_total_monthly_change = Portfolio_section_dict['Portfolio_total_monthly_change']
    Portfolio_started_c_monthly_change = Portfolio_section_dict['Portfolio_started_c_monthly_change']
    Portfolio_completed_c_monthly_change = Portfolio_section_dict['Portfolio_completed_c_monthly_change']
    Portfolio_total_yearly_change = Portfolio_section_dict['Portfolio_total_yearly_change']
    Portfolio_started_c_yearly_change = Portfolio_section_dict['Portfolio_started_c_yearly_change']
    Portfolio_completed_c_yearly_change = Portfolio_section_dict['Portfolio_completed_c_yearly_change']
    Portfolio_total_dwellings = Portfolio_section_dict['Portfolio_total_dwellings']
    Portfolio_completed_dwellings = Portfolio_section_dict['Portfolio_completed_dwellings']
    Portfolio_started_dwellings = Portfolio_section_dict['Portfolio_started_dwellings']
    Portfolio_in_programme_dwellings = Portfolio_section_dict['Portfolio_in_programme_dwellings']
    Portfolio_11_18m_started_c_pct = Portfolio_section_dict['Portfolio_11_18m_started_c_pct']
    Portfolio_18m_started_c_pct = Portfolio_section_dict['Portfolio_18m_started_c_pct']
    Portfolio_social_started_c_pct = Portfolio_section_dict['Portfolio_social_started_c_pct']
    Portfolio_private_started_c_pct = Portfolio_section_dict['Portfolio_private_started_c_pct']

    Portfolio_unknown_no = Portfolio_section_dict['Portfolio_unknown_no']

    # Section Title 
    paragraph = DR.add_paragraph('Overall remediation progress', style = 'Heading 2')

    # Figure Title
    paragraph = DR.add_paragraph(style = 'Normal')
    text = f'Figure {figure_count}: {Portfolio_started_c_no} residential buildings ({Portfolio_started_c_pct} of identified buildings) have started or completed remediation on unsafe cladding, of which {Portfolio_completed_c_no} ({Portfolio_completed_c_pct} of identified buildings) have completed remediation works.'
    run = paragraph.add_run(text)
    run.bold = True

    # Figure 
    DR.add_picture(os.path.join(paths_variables['figure_path'], f'Figure{figure_count}.svg'), width=Cm(17))
    figure_count += 1


    # Table caption
    text = f'Table {table_count}: Remediation progress for 11m+ buildings identified with unsafe cladding which are in the ACM programme, eligible for the Building Safety Fund and the CSS, remediated under the developer remediation contract and self-funded by registered providers of social housing, England, {cutoff}.'
    paragraph = DR.add_paragraph(style = 'Normal')
    run = paragraph.add_run(text)
    run.bold = True
    table_count += 1

    # Table
    table_data = Portfolio_remediation_table
    table_widths = [Cm(6.5), Cm(2.65), Cm(2.65), Cm(2.75), Cm(3.4)]
    table_heights = [Cm(1.15), Cm(1.75), Cm(0.55), Cm(0.55), Cm(0.55)]
    table = create_table(DR, table_data, table_widths, table_heights)

    # Table footer
    text = f'The in programme figures currently include {Portfolio_unknown_no} social self-funded buildings where their remediation status is unknown as registered providers are yet to provide remediation dates. We expect buildings with unknown remediation status to be confirmed in coming months as further data is collected through engagement with Homes England.'
    DR.add_paragraph(text, style = 'Normal')


    # Heading
    paragraph = DR.add_paragraph('Overall remediation: key statistics', style = 'Heading 3')

    # Paragraph
    text = f'Of the {Portfolio_total} residential buildings 11 metres and over in height with unsafe cladding the department is monitoring, as of {cutoff}:'
    DR.add_paragraph(text, style = 'Normal')

    # Bullet point
    text = f'{Portfolio_completed_c_no} buildings ({Portfolio_completed_c_pct}) have completed remediation, including those awaiting building control sign off'
    DR.add_paragraph(text, style = 'List Bullet')

    # Bullet point
    text = f'{Portfolio_started_nc_no} buildings ({Portfolio_started_nc_pct}) have started remediation'
    DR.add_paragraph(text, style = 'List Bullet')

    # Bullet point
    text = f'{Portfolio_in_programme_nc_no} buildings ({Portfolio_in_programme_nc_pct}) have not started remediation'
    DR.add_paragraph(text, style = 'List Bullet')

    # Paragraph
    text = f'Since the end of {last_month} {last_month_year}:'
    DR.add_paragraph(text, style = 'Normal')

    # Bullet point
    text = f'The department is monitoring the remediation progress of {Portfolio_total_monthly_change} buildings.'
    DR.add_paragraph(text, style = 'List Bullet')

    # Bullet point
    text = f'{Portfolio_started_c_monthly_change.capitalize()} buildings are known to have started or completed remediation, and {Portfolio_completed_c_monthly_change} buildings are known to have completed remediation.'
    DR.add_paragraph(text, style = 'List Bullet')


    # Paragraph
    text = f'There are an estimated {Portfolio_total_dwellings} dwellings in the occupied private and social sector 11m+ residential buildings with unsafe cladding that the department are monitoring. '
    text += f'Of these an estimated {Portfolio_completed_dwellings} dwellings are in buildings that have completed remediation, and an estimated {Portfolio_started_dwellings} additional dwellings are in buildings that have started remediation. An estimated {Portfolio_in_programme_dwellings} dwellings are in buildings that have not started remediation.'
    DR.add_paragraph(text, style = 'Normal')

    # Figure Title
    paragraph = DR.add_paragraph(style = 'Normal')
    text = f'Figure {figure_count}: Progress of remediating unsafe cladding differs across the programmes and monitoring routes due to the differing maturity of the programmes.'
    run = paragraph.add_run(text)
    run.bold = True
    
    # Figure
    DR.add_picture(os.path.join(paths_variables['figure_path'], f'Figure{figure_count}.svg'), width=Cm(17))
    figure_count += 1

    # Heading
    paragraph = DR.add_paragraph('Overall remediation by height', style = 'Heading 3')

    # Figure Title
    paragraph = DR.add_paragraph(style = 'Normal')
    text = f'Figure {figure_count}: {Portfolio_18m_started_c_pct} of the 18m+ buildings the department is monitoring the remediation progress of have started or completed remediation on unsafe cladding, compared to {Portfolio_11_18m_started_c_pct} of 11-18m buildings.'
    run = paragraph.add_run(text)
    run.bold = True
    

    # Figure 4
    DR.add_picture(os.path.join(paths_variables['figure_path'], f'Figure{figure_count}.svg'), width=Cm(17))
    figure_count += 1

    # Heading
    paragraph = DR.add_paragraph('Overall remediation by tenure', style = 'Heading 3')

    # Figure Title
    paragraph = DR.add_paragraph(style = 'Normal')
    text = f'Figure {figure_count}: {Portfolio_social_started_c_pct} of the social buildings the department is monitoring the remediation progress of have started or completed remediation on unsafe cladding, compared to {Portfolio_private_started_c_pct} of the private buildings.'
    run = paragraph.add_run(text)
    run.bold = True

    # Figure
    DR.add_picture(os.path.join(paths_variables['figure_path'], f'Figure{figure_count}.svg'), width=Cm(17))
    figure_count += 1

    # Figure caption
    DR.add_paragraph("The ‘Other’ bar includes high-rise buildings with unsafe ACM that are hotels, student accommodation and public buildings.", style = 'Normal')

    # Heading
    paragraph = DR.add_paragraph('Overall remediation by location', style = 'Heading 3')

    # Figure Title
    paragraph = DR.add_paragraph(style = 'Normal')
    text = f'Figure {figure_count}: Most buildings that the department are monitoring the cladding remediation of are concentrated around urbanised areas in England, particularly the urban areas of Greater London, Greater Manchester, West Yorkshire and the West of England.'
    run = paragraph.add_run(text)
    run.bold = True
    

    # Map heading
    text = f'England, {cutoff}'
    paragraph = DR.add_paragraph(text, style = 'Heading 3')

    # Map
    DR.add_picture(placeholder_path, width=Cm(17))
    figure_count += 1
    
    # Map caption
    DR.add_paragraph('Local authorities with 10 or fewer 11m+ buildings monitored with unsafe cladding are excluded from this map as their inclusion could lead to the identification of buildings with unsafe cladding.', style = 'Normal')
    
    return figure_count, table_count